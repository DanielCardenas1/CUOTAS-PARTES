import streamlit as st
from datetime import date
from decimal import Decimal, InvalidOperation
from sqlalchemy import text
import os

st.set_page_config(page_title="Cuotas Partes", page_icon="📑", layout="wide")

# --- Estilos institucionales ---
st.markdown("""
    <style>
    .main {background-color: #232a34; color: #e3e6ea;}
    .stButton > button {background-color: #1a2940; color: #fff; font-weight: bold; border-radius: 8px;}
    .stTable {background-color: #232a34;}
    .stTabs [data-baseweb=\"tab\"] {background-color: #1a2940; color: #fff;}
    </style>
""", unsafe_allow_html=True)

# --- Menú lateral principal ---
st.sidebar.title("Menú")
menu = st.sidebar.radio(
    "Navegación",
    (
        "🏠 Dashboard",
        "👤 Pensionados",
        "📑 Liquidaciones",
        "💰 Pagos",
        "📤 Cobro Persuasivo",
        "📊 Reportes y Seguimiento",
        "⚖️ Liquidaciones Masivas (30 Cuentas)",
        "🔒 Seguridad y Trazabilidad",
        "🗓️ Liquidar por Periodos Personalizados",
    ),
    index=0,
    key="menu_principal",
)


# --- Utilidades: número a letras (es-CO) ---
def numero_a_letras(n: int) -> str:
    """Convierte un número entero a texto en español (simplificado para valores grandes)."""
    unidades = (
        "cero", "uno", "dos", "tres", "cuatro", "cinco", "seis", "siete", "ocho", "nueve",
        "diez", "once", "doce", "trece", "catorce", "quince", "dieciséis", "diecisiete", "dieciocho", "diecinueve",
        "veinte", "veintiuno", "veintidós", "veintitrés", "veinticuatro", "veinticinco", "veintiséis", "veintisiete", "veintiocho", "veintinueve"
    )
    # Debe tener 10 entradas (índices 0..9) para evitar IndexError cuando d=9 (noventa)
    # Notar que 10-29 se manejan en el bloque num < 30, pero dejamos 'diez' y 'veinte' por completitud
    decenas = ("", "diez", "veinte", "treinta", "cuarenta", "cincuenta", "sesenta", "setenta", "ochenta", "noventa")
    centenas = ("", "cien", "doscientos", "trescientos", "cuatrocientos", "quinientos", "seiscientos", "setecientos", "ochocientos", "novecientos")

    def _decenas(num):
        if num < 30:
            return unidades[num]
        d, u = divmod(num, 10)
        if u == 0:
            return decenas[d]
        return f"{decenas[d]} y {unidades[u]}"

    def _centenas(num):
        if num < 100:
            return _decenas(num)
        c, r = divmod(num, 100)
        if c == 1:
            if r == 0:
                return "cien"
            return f"ciento {_decenas(r)}"
        return f"{centenas[c]}" if r == 0 else f"{centenas[c]} {_decenas(r)}"

    def _seccion(num, divisor, singular, plural):
        c, r = divmod(num, divisor)
        if c == 0:
            return "", r
        if c == 1:
            return f"{singular}", r
        return f"{numero_a_letras(c)} {plural}", r

    if n == 0:
        return "cero"
    if n < 0:
        return "menos " + numero_a_letras(-n)

    resultado = []
    millones, resto = _seccion(n, 1_000_000, "un millón", "millones")
    if millones:
        resultado.append(millones)
    miles, resto = _seccion(resto, 1_000, "mil", "mil")
    if miles:
        resultado.append(miles)
    if resto:
        resultado.append(_centenas(resto))
    return " ".join([p for p in resultado if p])


def generar_pdf_consolidado_en_memoria(entidad_nit: str, entidad_nombre: str, todas_las_cuentas: list, fecha_corte: date):
    """
    Genera el PDF consolidado EXACTAMENTE con la misma lógica/formatos del botón "📄 PDF Consolidado".
    Devuelve una tupla (bytes_pdf, nombre_archivo).
    """
    # Imports locales para evitar dependencias globales
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Image as RLImage, KeepTogether, Flowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    import io, calendar
    from dateutil.relativedelta import relativedelta

    buffer = io.BytesIO()
    # Margenes ligeramente menores para ganar ancho útil
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=0.25*inch, leftMargin=0.25*inch,
                            topMargin=0.35*inch, bottomMargin=0.35*inch)  # Márgenes más pequeños

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=14, spaceAfter=6, alignment=TA_CENTER, fontName='Helvetica-Bold')
    # Títulos de entidad (deudor/adeudado) un poco más pequeños para compactar
    entity_title_style = ParagraphStyle('EntityTitleSmall', parent=styles['Heading1'], fontSize=12, spaceAfter=4, alignment=TA_CENTER, fontName='Helvetica-Bold')
    normal_style = styles['Normal']
    normal_left_bold = ParagraphStyle('NormalLeftBold', parent=styles['Normal'], fontSize=10, alignment=TA_LEFT, fontName='Helvetica-Bold')

    story = []

    # Encabezado: ciudad y CUENTA DE COBRO + No.
    story.append(Spacer(1, 0.25*inch))  # Menos espacio arriba
    header_line1 = Table([["BOGOTÁ, D.C.", "CUENTA DE COBRO"]], colWidths=[4*inch, 3*inch])
    header_line1.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),  # Título más pequeño
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    story.append(header_line1)

    # Número de cuenta (global, único) desde BD: registrar/usar CONSOLIDADO para la entidad y periodo
    from app.db import get_session as _get_session
    from app.models import CuentaCobro as _CuentaCobro
    from sqlalchemy import select as _select, func as _func
    from datetime import date as _date, datetime as _datetime

    from dateutil.relativedelta import relativedelta
    _periodo_fin = _date(fecha_corte.year, fecha_corte.month, 1)
    _periodo_ini = _periodo_fin - relativedelta(months=29)

    consecutivo_visible = None
    ahora = _date.today()
    with _get_session() as _s:
        existente = _s.execute(
            _select(_CuentaCobro).where(
                _CuentaCobro.nit_entidad == str(entidad_nit),
                _CuentaCobro.pensionado_identificacion == '__CONSOLIDADO__',
                _CuentaCobro.periodo_inicio == _periodo_ini,
                _CuentaCobro.periodo_fin == _periodo_fin,
            ).order_by(_CuentaCobro.fecha_creacion.desc())
        ).scalars().first()
        if existente is not None:
            consecutivo_visible = existente.consecutivo
        else:
            # Siguiente global: MAX(consecutivo)+1
            mx = _s.execute(_select(_func.max(_CuentaCobro.consecutivo))).scalar()
            consecutivo_visible = (mx or 0) + 1
            # Crear registro CONSOLIDADO
            reg = _CuentaCobro(
                consecutivo=consecutivo_visible,
                nit_entidad=str(entidad_nit),
                empresa=entidad_nombre,
                pensionado_identificacion='__CONSOLIDADO__',
                pensionado_nombre=f'CONSOLIDADO {entidad_nombre}',
                periodo_inicio=_periodo_ini,
                periodo_fin=_periodo_fin,
                total_capital=None,
                total_intereses=None,
                total_liquidacion=None,
                archivo_pdf=None,
                estado='CONSOLIDADO',
                version=1,
                fecha_creacion=_datetime.now(),
                fecha_actualizacion=_datetime.now(),
            )
            _s.add(reg)
            _s.commit()

    numero_table = Table([["", f"No.  {consecutivo_visible}"]], colWidths=[4*inch, 3*inch])
    numero_table.setStyle(TableStyle([
        ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 0), (1, 0), 12),  # Número más pequeño
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
    ]))
    story.append(numero_table)

    story.append(Spacer(1, 0.2*inch))

    # Entidad acreedora (encabezado)
    entidad_nombre_pdf = entidad_nombre
    story.append(Paragraph(f"<b>{entidad_nombre_pdf.upper()}</b>", entity_title_style))
    story.append(Paragraph(f"<b>NIT: {entidad_nit}</b>", entity_title_style))

    story.append(Spacer(1, 0.25*inch))

    # DEBE A: (siempre SENA)
    story.append(Paragraph("<b>DEBE A:</b>", title_style))
    story.append(Spacer(1, 0.2*inch))
    deudor_nombre_const = "SERVICIO NACIONAL DE APRENDIZAJE - SENA"
    deudor_nit_const = "899,999,034-1"
    story.append(Paragraph(f"<b>{deudor_nombre_const}</b>", entity_title_style))
    story.append(Paragraph(f"<b>{deudor_nit_const}</b>", entity_title_style))

    # Totales consolidados (usando la misma suma que el botón)
    total_consolidado_capital = 0.0
    total_consolidado_intereses = 0.0
    total_consolidado_total = 0.0
    for p in todas_las_cuentas:
        # Soportar estructura mínima (sin totales por cuenta)
        capital_pensionado = sum(float(cuenta.get('capital_total', 0.0)) for cuenta in p.get('cuentas', []))
        intereses_pensionado = sum(float(cuenta.get('intereses', 0.0)) for cuenta in p.get('cuentas', []))
        total_pensionado = capital_pensionado + intereses_pensionado
        total_consolidado_capital += capital_pensionado
        total_consolidado_intereses += intereses_pensionado
        total_consolidado_total += total_pensionado

    story.append(Spacer(1, 0.2*inch))

    # Valor en letras y numérico
    valor_letras = numero_a_letras(int(total_consolidado_total))
    story.append(Paragraph(f"<b>LA SUMA DE: {valor_letras.upper()} PESOS M/CTE</b>", normal_style))
    story.append(Paragraph(f"<b>$ {total_consolidado_total:,.0f}</b>", normal_left_bold))

    story.append(Spacer(1, 0.25*inch))

    # Concepto con periodo de 30 meses (inicio-fin), meses en español
    meses_es = {1:'enero',2:'febrero',3:'marzo',4:'abril',5:'mayo',6:'junio',7:'julio',8:'agosto',9:'septiembre',10:'octubre',11:'noviembre',12:'diciembre'}
    _fecha_fin = date(fecha_corte.year, fecha_corte.month, 1)
    _fecha_ini = _fecha_fin - relativedelta(months=29)
    inicio_txt = f"01 de {meses_es[_fecha_ini.month].capitalize()} de {_fecha_ini.year}"
    ultimo_dia = calendar.monthrange(_fecha_fin.year, _fecha_fin.month)[1]
    fin_txt = f"{ultimo_dia} de {meses_es[_fecha_fin.month].capitalize()} de {_fecha_fin.year}"
    concepto_text = (
        "Por concepto de Cuotas Partes Pensionales sobre pagos realizados en la nómina de Pensionados del SENA "
        f"a fecha de corte {inicio_txt} a corte de {fin_txt} por los pensionados que se relacionan a continuación:"
    )
    story.append(Paragraph(concepto_text, normal_style))
    story.append(Spacer(1, 0.2*inch))

    # Tabla principal
    tabla_data = [[
        'No. Cédula', 'Apellidos y Nombres', 'Ingreso\nNómina', 'RESOLUCIÓN\nNº', '% Cuota\nParte', 'Base\nCálculo\nCuota',
        'Saldo\nCapital\nCausado', 'Intereses\nAcumulados', 'TOTAL\nDEUDA'
    ]]

    for p in todas_las_cuentas:
        capital_pensionado = sum(float(cuenta.get('capital_total', 0.0)) for cuenta in p.get('cuentas', []))
        intereses_pensionado = sum(float(cuenta.get('intereses', 0.0)) for cuenta in p.get('cuentas', []))
        total_pensionado = capital_pensionado + intereses_pensionado
        base_calc = float(p['pensionado'].get('base_calculo_cuota', 0.0))
        nombre = p['pensionado']['nombre']
        nombre_corto = nombre[:25] + '...' if len(nombre) > 25 else nombre
        # Formateo de ingreso nómina y resolución desde la data del pensionado
        ingreso_nomina = p['pensionado'].get('ingreso_nomina')
        ingreso_txt = ''
        try:
            if ingreso_nomina and hasattr(ingreso_nomina, 'strftime'):
                meses_abbr = ['ene', 'feb', 'mar', 'abr', 'may', 'jun', 'jul', 'ago', 'sep', 'oct', 'nov', 'dic']
                d = ingreso_nomina.day
                m = meses_abbr[ingreso_nomina.month - 1]
                y = str(ingreso_nomina.year)[-2:]
                ingreso_txt = f"{d}-{m}-{y}"
            else:
                ingreso_txt = str(ingreso_nomina or '')
        except Exception:
            ingreso_txt = str(ingreso_nomina or '')

        resol_txt = str(p['pensionado'].get('resolucion') or '')

        tabla_data.append([
            str(p['pensionado']['cedula']),
            nombre_corto,
            ingreso_txt or '1-may-08',
            resol_txt or '3089 de 2007',
            f"{float(p['pensionado']['porcentaje_cuota'])*100:.2f}%",
            f"{base_calc:,.0f}",
            f"{capital_pensionado:,.0f}",
            f"{intereses_pensionado:,.0f}",
            f"{total_pensionado:,.0f}"
        ])

    # Totales fila final
    tabla_data.append([
        '', '', '', 'TOTAL', '', '',
        f"{total_consolidado_capital:,.0f}",
        f"{total_consolidado_intereses:,.0f}",
        f"{total_consolidado_total:,.0f}"
    ])

    # Ajuste de anchos: usar el ancho útil de la página casi total para evitar recortes de valores grandes
    try:
        available_width = doc.width  # ancho útil (página - márgenes)
    except Exception:
        from reportlab.lib.pagesizes import letter as _letter
        available_width = _letter[0] - (0.3*inch + 0.3*inch)
    # Dejar solo ~0.1in de aire para maximizar el ancho de columnas
    target_width = max(6.2*inch, available_width - 0.1*inch)

    # Pesos relativos por columna, ampliando nombres y columnas numéricas para montos altos
    col_weights = [1.0, 2.6, 0.9, 1.0, 0.7, 1.0, 1.1, 1.1, 1.2]
    weights_sum = sum(col_weights)
    col_widths = [w/weights_sum * target_width for w in col_weights]

    tabla_pensionados = Table(tabla_data, colWidths=col_widths, hAlign='LEFT')
    light_header = colors.Color(0.94, 0.94, 0.94)
    light_row = colors.Color(0.985, 0.985, 0.985)
    tabla_pensionados.setStyle(TableStyle([
        # Encabezado gris claro, tipografía consistente y centrado
        ('BACKGROUND', (0, 0), (-1, 0), light_header),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 7.5),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),

        # Filas: tamaño discreto, alineaciones refinadas como el ejemplo
        ('FONTNAME', (0, 1), (-1, -2), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -2), 7.5),
        ('ALIGN', (0, 1), (0, -2), 'RIGHT'),   # Cédula
        ('ALIGN', (1, 1), (1, -2), 'LEFT'),    # Nombres
        ('ALIGN', (2, 1), (3, -2), 'CENTER'),  # Ingreso nómina y Resolución
        ('ALIGN', (4, 1), (-1, -2), 'RIGHT'),  # % y valores
        ('VALIGN', (0, 1), (-1, -2), 'MIDDLE'),
        # Alternancia muy sutil para elegancia (casi blanco)
        ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, light_row]),

        # Fila de totales (última): énfasis en negrita y línea superior separadora
        ('BACKGROUND', (0, -1), (-1, -1), colors.white),
        ('FONTNAME', (5, -1), (5, -1), 'Helvetica-Bold'),
        ('FONTNAME', (6, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (5, -1), (-1, -1), 9),
        ('ALIGN', (6, -1), (-1, -1), 'RIGHT'),
        ('LINEABOVE', (0, -1), (-1, -1), 1, colors.black),

        # Grid delgado y paddings ajustados
        ('GRID', (0, 0), (-1, -1), 0.8, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 1.5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 1.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 1.5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 1.5),
    ]))
    story.append(tabla_pensionados)

    # Notas de pago
    story.append(Spacer(1, 0.3*inch))
    nota_style = ParagraphStyle('NotaConsol', parent=styles['Normal'], fontSize=9, alignment=TA_LEFT, fontName='Helvetica')

    story.append(Paragraph("El pago de la presente cuenta se debe hacer a traves de nuestro <b>SISTEMA DE PAGOS EN LINEA - PSE</b>,", nota_style))
    story.append(Spacer(1, 0.1*inch))
    story.append(Paragraph(
        "<u>Las obligaciones generadas con posterioridad al 29 de julio de 2006 por concepto de cuotas partes pensionales causaran un interes</u>", nota_style
    ))
    story.append(Paragraph(
        "<u>del DTF entre la fecha de pago de la mesada pensional y la fecha de reembolso por parte de la entidad concurrente.</u>", nota_style
    ))
    story.append(Paragraph(
        "(Articulo 4 Ley 1066 de 2006,Circular N. 1-2006. 3-2006-016603).", nota_style
    ))
    # Bloque de la Directora centrado
    firma_center_name = ParagraphStyle('FirmaCenterNameConsol', parent=styles['Normal'], fontSize=11, alignment=TA_CENTER, fontName='Helvetica-Bold')
    firma_center_sub = ParagraphStyle('FirmaCenterSubConsol', parent=styles['Normal'], fontSize=10, alignment=TA_CENTER, fontName='Helvetica')
    story.append(Spacer(1, 0.8*inch))
    story.append(Paragraph("ADRIANA MILENA GASCA CARDOSO", firma_center_name))
    story.append(Paragraph("Directora Administrativa y Financiera", firma_center_sub))
    story.append(Paragraph("Sena - Dirección General", firma_center_sub))

    # Imagen de firmas (una sola) alineada a la izquierda
    import os
    from reportlab.lib.utils import ImageReader
    firmas_dir = os.path.join(os.path.dirname(__file__), 'Firmas')
    candidatos = [
        'Captura de pantalla 2025-10-04 111523.png',
        'firmas.png',
        'firmas.jpg',
        'Firmas.png',
        'Firmas.jpg',
    ]
    firma_path = None
    for cand in candidatos:
        pth = os.path.join(firmas_dir, cand)
        if os.path.exists(pth):
            firma_path = pth
            break

    if firma_path:
        class LeftSignatureImage(Flowable):
            def __init__(self, path: str, max_width=9.0*cm, max_height=4.0*cm):
                super().__init__()
                self.path = path
                self.max_width = max_width
                self.max_height = max_height
                self._w = max_width
                self._h = max_height

            def wrap(self, availWidth, availHeight):
                try:
                    ir = ImageReader(self.path)
                    iw, ih = ir.getSize()
                    # Escalar para que quepa en los máximos conservando proporción
                    sx = self.max_width / float(iw)
                    sy = self.max_height / float(ih)
                    s = min(sx, sy)
                    self._w = iw * s
                    self._h = ih * s
                except Exception:
                    self._w = min(self.max_width, availWidth)
                    self._h = self.max_height
                return availWidth, self._h

            def draw(self):
                try:
                    self.canv.drawImage(self.path, 0, 0, width=self._w, height=self._h, mask='auto')
                except Exception:
                    pass

        story.append(Spacer(1, 0.4*inch))
        story.append(LeftSignatureImage(firma_path))

    # Anexo por pensionado con detalle de las 30 cuentas
    story.append(PageBreak())
    subtitle_pensionado = ParagraphStyle('SubPens', parent=styles['Heading2'], fontSize=12, spaceAfter=4, alignment=TA_LEFT, fontName='Helvetica-Bold')
    small_right = ParagraphStyle('SmallRight', parent=styles['Normal'], fontSize=8, alignment=TA_RIGHT)
    small_center = ParagraphStyle('SmallCenter', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER)
    small_left = ParagraphStyle('SmallLeft', parent=styles['Normal'], fontSize=8, alignment=TA_LEFT)

    for idx_p, p in enumerate(todas_las_cuentas):
        if idx_p > 0:
            story.append(PageBreak())
        # Definir un ancho de caja único para todas las tablas del anexo y centrar
        box_bleed = 1.0*cm
        box_width = max(6.6*inch, doc.width - box_bleed)

        # Cabecera de anexo: título a la izquierda y 'Liquidación Nro. <consecutivo>' a la derecha
        header_left = Paragraph(
            f"Detalle de cuentas de cobro - {p['pensionado']['nombre']} (CC {p['pensionado']['cedula']})",
            subtitle_pensionado
        )
        header_right = Paragraph(
            f"Liquidación Nro. {consecutivo_visible}",
            ParagraphStyle('SubPensRight', parent=subtitle_pensionado, alignment=TA_RIGHT)
        )
        header_tbl = Table([[header_left, header_right]], colWidths=[box_width*0.65, box_width*0.35], hAlign='CENTER')
        header_tbl.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(header_tbl)

        # Tabla de títulos como en la liquidación individual (Entidad/Período/Nombre/Identificación)
        try:
            nit_text = str(entidad_nit)
            # Periodo en formato 'ABRIL DEL 2023 A AGOSTO DEL 2025'
            periodo_texto = f"{meses_es[_fecha_ini.month].upper()} DEL {_fecha_ini.year} A {meses_es[_fecha_fin.month].upper()} DEL {_fecha_fin.year}"
            info_superior = [
                ['Entidad', f"{entidad_nombre} - NIT. {nit_text}"],
                ['Período', f"CUOTAS PARTES POR COBRAR {periodo_texto}"],
                ['Nombre', p['pensionado']['nombre']],
                ['Identificación', str(p['pensionado']['cedula'])],
            ]
            tabla_superior = Table(info_superior, colWidths=[3*cm, box_width - 3*cm], hAlign='CENTER')
            tabla_superior.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(tabla_superior)
            story.append(Spacer(1, 0.2*cm))
        except Exception:
            pass

        # Tabla de resumen intermedia (Ingreso a Nómina, % Cuota Parte, Estado, Capital, Interes, Total)
        try:
            # Ingreso a Nómina formateado tipo dd/mm/aaaa
            ingreso_nomina = p['pensionado'].get('ingreso_nomina')
            ingreso_txt_res = ''
            if ingreso_nomina and hasattr(ingreso_nomina, 'strftime'):
                ingreso_txt_res = ingreso_nomina.strftime('%d/%m/%Y')
            else:
                ingreso_txt_res = str(ingreso_nomina or '')

            porcentaje_txt = f"{float(p['pensionado'].get('porcentaje_cuota', 0))*100:.2f}%"

            # Calcular subtotales del periodo (para Capital pendiente/Interes/Total del resumen)
            cuentas_ord = sorted(p['cuentas'], key=lambda c: (c['año'], c['mes']))
            subtotal_capital = subtotal_intereses = subtotal_total = 0.0
            for cta in cuentas_ord:
                cap = float(cta.get('capital_total', 0))
                inte = float(cta.get('intereses', 0))
                tot = float(cta.get('total_cuenta', cap + inte))
                subtotal_capital += cap
                subtotal_intereses += inte
                subtotal_total += tot

            resumen_data = [
                ['Ingreso a Nómina', '% Cuota Parte', 'Estado', 'Capital pendiente', 'Interes acumulado', 'Total'],
                [ingreso_txt_res, porcentaje_txt, 'ACTIVO', f"${subtotal_capital:,.2f}", f"${subtotal_intereses:,.2f}", f"${subtotal_total:,.2f}"],
            ]
            # Escalar las columnas para que el cuadro sea un poco más angosto y centrado
            resumen_weights_cm = [2.3, 2.3, 2.3, 2.3, 2.4, 2.4]
            total_weights_pts = sum(resumen_weights_cm) * cm
            scale = box_width / float(total_weights_pts) if total_weights_pts else 1.0
            resumen_colwidths = [(w*cm)*scale for w in resumen_weights_cm]
            tabla_resumen = Table(resumen_data, colWidths=resumen_colwidths, hAlign='CENTER')
            tabla_resumen.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('BACKGROUND', (2, 1), (2, 1), colors.green),  # Estado en verde
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 6),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 2),
                ('RIGHTPADDING', (0, 0), (-1, -1), 2),
                ('TOPPADDING', (0, 0), (-1, -1), 1),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
            ]))
            story.append(tabla_resumen)
            story.append(Spacer(1, 0.2*inch))
        except Exception:
            pass

        detalle_headers = [
            Paragraph('#', small_center),
            Paragraph('Período', small_center),
            Paragraph('Capital Total', small_center),
            Paragraph('Intereses', small_center),
            Paragraph('Total Cuenta', small_center),
        ]
        detalle_data = [detalle_headers]

        cuentas_ord = sorted(p['cuentas'], key=lambda c: (c['año'], c['mes']))
        subtotal_capital = subtotal_intereses = subtotal_total = 0.0
        for idx, cta in enumerate(cuentas_ord, start=1):
            periodo_txt = f"{meses_es[cta['mes']].upper()} {cta['año']}"
            cap = float(cta.get('capital_total', 0))
            inte = float(cta.get('intereses', 0))
            tot = float(cta.get('total_cuenta', cap + inte))
            subtotal_capital += cap
            subtotal_intereses += inte
            subtotal_total += tot
            detalle_data.append([
                Paragraph(f"{idx}", small_center),
                Paragraph(periodo_txt, small_center),
                Paragraph(f"${cap:,.2f}", small_right),
                Paragraph(f"${inte:,.2f}", small_right),
                Paragraph(f"${tot:,.2f}", small_right),
            ])

        # Ya no agregamos una fila 'TOTAL' dentro de esta tabla; pondremos un resumen separado como en el ejemplo

        # Escalar las columnas del detalle para usar el mismo box_width
        detalle_weights_in = [0.5, 1.3, 1.35, 1.25, 1.35]
        total_weights_pts = sum(detalle_weights_in) * inch
        scale_det = box_width / float(total_weights_pts) if total_weights_pts else 1.0
        detalle_colwidths = [(w*inch)*scale_det for w in detalle_weights_in]
        detalle_table = Table(detalle_data, colWidths=detalle_colwidths, repeatRows=1, hAlign='CENTER')
        detalle_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTSIZE', (0, 0), (-1, 0), 7.5),  # encabezado detalle
            ('FONTSIZE', (0, 1), (-1, -1), 7.5),  # filas
            ('TOPPADDING', (0, 0), (-1, -1), 1.5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 1.5),
        ]))
        story.append(detalle_table)

        # Resumen final de totales (3 renglones) como la segunda imagen
        try:
            # Recalcular subtotales (o reutilizarlos si se desea)
            cuentas_ord = sorted(p['cuentas'], key=lambda c: (c['año'], c['mes']))
            subtotal_capital = subtotal_intereses = subtotal_total = 0.0
            for cta in cuentas_ord:
                cap = float(cta.get('capital_total', 0))
                inte = float(cta.get('intereses', 0))
                tot = float(cta.get('total_cuenta', cap + inte))
                subtotal_capital += cap
                subtotal_intereses += inte
                subtotal_total += tot

            ini_str = f"{meses_es[_fecha_ini.month]} {_fecha_ini.year}"
            fin_str = f"{meses_es[_fecha_fin.month]} de {_fecha_fin.year}"

            totales_data = [
                [f"Total capital adeudado de {ini_str} a {fin_str}", f"${subtotal_capital:,.0f}"],
                [f"Intereses causados (Ley 1066 de 2006) de {ini_str} a {fin_str}", f"${subtotal_intereses:,.0f}"],
                [f"Total de la deuda con corte {fin_str}", f"${subtotal_total:,.0f}"],
            ]

            # Ancho del cuadro de totales: mismo 'box_width' que arriba y centrado
            amount_w = 4.2*cm
            tot_col_widths = [box_width - amount_w, amount_w]
            small_bold = ParagraphStyle('SmallBold', parent=small_left, fontName='Helvetica-Bold')
            tot_tbl = Table(
                [[Paragraph(t[0], small_left), Paragraph(t[1], small_right)] for t in totales_data],
                colWidths=tot_col_widths,
                hAlign='CENTER'
            )
            tot_tbl.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
            ]))
            story.append(Spacer(1, 0.15*inch))
            story.append(tot_tbl)
        except Exception:
            pass

        story.append(Spacer(1, 0.25*inch))

    # Construir PDF y retornar bytes + nombre
    doc.build(story)
    pdf_data = buffer.getvalue()
    buffer.close()
    pdf_name = f"LIQUIDACION_CONSOLIDADA_{entidad_nit}_{fecha_corte.strftime('%Y%m%d')}.pdf"
    return pdf_data, pdf_name, consecutivo_visible


def unir_pdfs_en_memoria(pdf1_bytes: bytes, pdf2_bytes: bytes) -> bytes:
    """Une dos PDFs en memoria tal cual están: pdf1 seguido de pdf2.
    No altera contenidos, sólo concatena páginas. Devuelve bytes del PDF resultante."""
    try:
        from pypdf import PdfReader, PdfWriter
        import io
        w = PdfWriter()
        r1 = PdfReader(io.BytesIO(pdf1_bytes))
        r2 = PdfReader(io.BytesIO(pdf2_bytes))
        for p in r1.pages:
            w.add_page(p)
        for p in r2.pages:
            w.add_page(p)
        out = io.BytesIO()
        w.write(out)
        w.close()
        return out.getvalue()
    except Exception as e:
        raise RuntimeError(f"No fue posible unir los PDFs: {e}")


def convert_docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    """Convierte DOCX (bytes) a PDF (bytes) en Windows de forma robusta.
    Estrategia: docx2pdf (CLI) -> python -m docx2pdf -> docx2pdf.convert con COM -> pywin32 directo.
    Incluye 2-3 reintentos para errores COM transitorios (call rejected)."""
    import tempfile, os, shutil, subprocess, sys, time
    tmpdir = tempfile.mkdtemp(prefix="carta_conv_")
    in_path = os.path.join(tmpdir, "doc.docx")
    out_path = os.path.join(tmpdir, "doc.pdf")
    try:
        with open(in_path, 'wb') as f:
            f.write(docx_bytes)

        # 1) CLI docx2pdf
        def try_cli() -> tuple[bool, str]:
            stderr = ''
            completed = None
            try:
                completed = subprocess.run(["docx2pdf", in_path, out_path], capture_output=True, text=True, timeout=180)
            except Exception:
                pass
            if not (completed and completed.returncode == 0 and os.path.exists(out_path)):
                try:
                    completed = subprocess.run([sys.executable, "-m", "docx2pdf", in_path, out_path], capture_output=True, text=True, timeout=180)
                except Exception:
                    completed = None
            ok = bool(completed and completed.returncode == 0 and os.path.exists(out_path))
            if completed and completed.stderr:
                stderr = completed.stderr
            return ok, stderr

        ok, cli_stderr = try_cli()
        if not ok:
            # 2) In-proc docx2pdf con COM (ctypes)
            try:
                from docx2pdf import convert as docx2pdf_convert
                import ctypes
                ole32 = ctypes.windll.ole32
                # Reintentos por call rejected
                for i in range(3):
                    try:
                        ole32.CoInitialize(None)
                        try:
                            docx2pdf_convert(in_path, out_path)
                        finally:
                            ole32.CoUninitialize()
                        break
                    except Exception as e:
                        if i == 2:
                            raise e
                        time.sleep(1.5)
            except Exception as e1:
                # 3) pywin32 directo
                try:
                    import pythoncom
                    pythoncom.CoInitialize()
                    try:
                        import win32com.client
                        from win32com.client import constants
                        # Reintentos por call rejected
                        for i in range(3):
                            try:
                                word = win32com.client.DispatchEx('Word.Application')
                                word.Visible = False
                                doc = word.Documents.Open(in_path, ReadOnly=True)
                                # 17 = wdExportFormatPDF
                                doc.ExportAsFixedFormat(out_path, 17)
                                doc.Close(False)
                                word.Quit()
                                break
                            except Exception as e2:
                                try:
                                    # intentar cerrar/quitar
                                    word.Quit()
                                except Exception:
                                    pass
                                if i == 2:
                                    raise e2
                                time.sleep(1.5)
                    finally:
                        pythoncom.CoUninitialize()
                except Exception as e2:
                    raise RuntimeError(f"Fallo conversión DOCX->PDF. CLI: {cli_stderr}\nerr1: {e1}\nerr2: {e2}")

        if not os.path.exists(out_path):
            raise RuntimeError("La conversión no generó el PDF esperado.")
        with open(out_path, 'rb') as f:
            return f.read()
    finally:
        try:
            shutil.rmtree(tmpdir, ignore_errors=True)
        except Exception:
            pass


def generar_carta_docx_asunto_en_memoria(template_path: str, entidad_nombre: str, fecha_corte: date, todas_las_cuentas: list | None = None):
    """
    Carga el .docx de modelo, reemplaza:
    - Asunto con el periodo y la entidad
    - Párrafo del cuerpo con total en letras/número (si hay consolidado)
    - ANEXOS (encabezado, item1, item2)
    - Firma (nombre y cargo)
    Devuelve (bytes_docx, nombre_archivo) con un DOCX válido para Microsoft Word.
    """
    import zipfile, io, calendar
    import xml.etree.ElementTree as ET

    # Calcular textos de fechas en español
    meses_es = {1:'enero',2:'febrero',3:'marzo',4:'abril',5:'mayo',6:'junio',7:'julio',8:'agosto',9:'septiembre',10:'octubre',11:'noviembre',12:'diciembre'}
    from dateutil.relativedelta import relativedelta
    _fecha_fin = date(fecha_corte.year, fecha_corte.month, 1)
    _fecha_ini = _fecha_fin - relativedelta(months=29)
    inicio_txt = f"01 de {meses_es[_fecha_ini.month]} de {_fecha_ini.year}"
    ultimo_dia = calendar.monthrange(_fecha_fin.year, _fecha_fin.month)[1]
    fin_txt = f"{ultimo_dia} de {meses_es[_fecha_fin.month]} de {_fecha_fin.year}"
    asunto_nuevo = (
        f"Asunto: Cuenta de cobro de cuotas partes pensionales del periodo comprendido entre {inicio_txt} al {fin_txt} – {entidad_nombre}"
    )

    # Calcular totales si se proporcionan las cuentas consolidadas
    total_consolidado_capital = 0.0
    total_consolidado_intereses = 0.0
    total_consolidado_total = 0.0
    if todas_las_cuentas:
        try:
            for p in todas_las_cuentas:
                cap = sum(float(c.get('capital_total', 0.0)) for c in p.get('cuentas', []))
                inte = sum(float(c.get('intereses', 0.0)) for c in p.get('cuentas', []))
                total_consolidado_capital += cap
                total_consolidado_intereses += inte
            total_consolidado_total = total_consolidado_capital + total_consolidado_intereses
        except Exception:
            pass

    valor_letras = numero_a_letras(int(total_consolidado_total)).upper() if total_consolidado_total else None

    # Opción 1: Usar python-docx si está disponible (más seguro para Word)
    try:
        from docx import Document
        import io as _io

        doc = Document(template_path)
        # Construir textos dependientes
        mes_ini_txt = f"{meses_es[_fecha_ini.month]} {_fecha_ini.year}"
        mes_fin_txt = f"{meses_es[_fecha_fin.month]} {_fecha_fin.year}"
        item1_txt = f"Cuenta de Cobro del periodo {mes_ini_txt} a {mes_fin_txt}."
        item2_txt = f"Liquidación Oficial de Pensionados de la {entidad_nombre}."
        if valor_letras is not None and total_consolidado_total:
            cuerpo_txt = (
                f"Me permito remitir la Cuenta de Cobro a cargo de esa Entidad, por concepto de Cuotas Partes Pensionales, por el periodo comprendido del {inicio_txt} al {fin_txt} el valor corresponde a {valor_letras} PESOS M/CTE. (${total_consolidado_total:,.0f})."
            )
        else:
            cuerpo_txt = (
                f"Me permito remitir la Cuenta de Cobro a cargo de esa Entidad, por concepto de Cuotas Partes Pensionales, por el periodo comprendido del {inicio_txt} al {fin_txt}."
            )

        def replace_paragraph_text(paragraph, new_text: str):
            # Reemplaza todo el texto del párrafo por new_text preservando el estilo del párrafo
            paragraph.clear() if hasattr(paragraph, 'clear') else None
            paragraph.text = new_text

        for p in doc.paragraphs:
            txt = p.text.strip()
            upper = txt.upper()
            if txt.startswith('Asunto:'):
                replace_paragraph_text(p, asunto_nuevo)
            elif txt.startswith('Me permito remitir la Cuenta de Cobro'):
                replace_paragraph_text(p, cuerpo_txt)
            elif upper.startswith('ANEXOS'):
                replace_paragraph_text(p, 'ANEXOS:')
            elif txt.startswith('Cuenta de Cobro del'):
                replace_paragraph_text(p, item1_txt)
            elif txt.startswith('Liquidación Oficial Pensionados del'):
                replace_paragraph_text(p, item2_txt)
            elif ('JOSE GIOVANNI' in upper) or ('LOZANO BOLIVAR' in upper):
                replace_paragraph_text(p, 'JOSE GIOVANNI LOZANO BOLIVAR')
            elif ('COORDINADOR' in upper) and ('CARTERA' in upper):
                replace_paragraph_text(p, 'Coordinador Grupo Recaudo y Cartera')

        out_buf = _io.BytesIO()
        doc.save(out_buf)
        out_bytes = out_buf.getvalue()
        out_buf.close()
        file_name = f"Carta_Cuenta_Cobro_{fecha_corte.strftime('%Y%m%d')}.docx"
        return out_bytes, file_name
    except Exception:
        # Opción 2: Fallback XML directo (mantener implementación existente)
        with zipfile.ZipFile(template_path, 'r') as zin:
            xml_bytes = zin.read('word/document.xml')
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        root = ET.fromstring(xml_bytes)

        asunto_reemplazado = False
        cuerpo_reemplazado = False
        item1_reemplazado = False
        item2_reemplazado = False
        anexos_reemplazado = False
        cordial_reemplazado = False
        firma_nombre_reemplazada = False
        firma_cargo_reemplazado = False
        for p in root.findall('.//w:p', ns):
            parts = []
            for t in p.findall('.//w:t', ns):
                if t.text:
                    parts.append(t.text)
            para_text = ''.join(parts).strip() if parts else ''
            if para_text.startswith('Asunto:'):
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                t = ET.SubElement(r, f"{{{ns['w']}}}t")
                t.text = asunto_nuevo
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                asunto_reemplazado = True
                continue

            if (not cuerpo_reemplazado) and para_text.startswith('Me permito remitir la Cuenta de Cobro'):
                if valor_letras is not None and total_consolidado_total:
                    cuerpo_txt = (
                        f"Me permito remitir la Cuenta de Cobro a cargo de esa Entidad, por concepto de Cuotas Partes Pensionales, por el periodo comprendido del {inicio_txt} al {fin_txt} el valor corresponde a {valor_letras} PESOS M/CTE. (${total_consolidado_total:,.0f})."
                    )
                else:
                    cuerpo_txt = (
                        f"Me permito remitir la Cuenta de Cobro a cargo de esa Entidad, por concepto de Cuotas Partes Pensionales, por el periodo comprendido del {inicio_txt} al {fin_txt}."
                    )
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = cuerpo_txt
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                cuerpo_reemplazado = True
                continue

            if (not anexos_reemplazado) and para_text.upper().startswith('ANEXOS'):
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = 'ANEXOS:'
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                anexos_reemplazado = True
                continue

            if (not item1_reemplazado) and para_text.startswith('Cuenta de Cobro del'):
                mes_ini_txt = f"{meses_es[_fecha_ini.month]} {_fecha_ini.year}"
                mes_fin_txt = f"{meses_es[_fecha_fin.month]} {_fecha_fin.year}"
                item1_txt = f"Cuenta de Cobro del periodo {mes_ini_txt} a {mes_fin_txt}."
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = item1_txt
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                item1_reemplazado = True
                continue

            if (not item2_reemplazado) and para_text.startswith('Liquidación Oficial Pensionados del'):
                item2_txt = f"Liquidación Oficial de Pensionados de la {entidad_nombre}."
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = item2_txt
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                item2_reemplazado = True
                continue

            if (not cordial_reemplazado) and para_text.lower().startswith('cordial'):
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = 'Cordialmente,'
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                cordial_reemplazado = True
                continue

            if (not firma_nombre_reemplazada) and ('JOSE GIOVANNI' in para_text.upper() or 'LOZANO BOLIVAR' in para_text.upper()):
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = 'JOSE GIOVANNI LOZANO BOLIVAR'
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                firma_nombre_reemplazada = True
                continue

            if (not firma_cargo_reemplazado) and ('COORDINADOR' in para_text.upper() and 'CARTERA' in para_text.upper()):
                to_remove = []
                for child in list(p):
                    if child.tag == f"{{{ns['w']}}}r":
                        to_remove.append(child)
                for r in to_remove:
                    p.remove(r)
                r = ET.SubElement(p, f"{{{ns['w']}}}r")
                tnode = ET.SubElement(r, f"{{{ns['w']}}}t")
                tnode.text = 'Coordinador Grupo Recaudo y Cartera'
                tnode.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                firma_cargo_reemplazado = True
                continue

        new_xml = ET.tostring(root, encoding='utf-8')
        out_buf = io.BytesIO()
        with zipfile.ZipFile(template_path, 'r') as zin, zipfile.ZipFile(out_buf, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = new_xml
                zout.writestr(item, data)
        out_bytes = out_buf.getvalue()
        out_buf.close()
        file_name = f"Carta_Cuenta_Cobro_{fecha_corte.strftime('%Y%m%d')}.docx"
        return out_bytes, file_name


def generar_readme_texto(entidad_nit: str, total_pensionados: int, total_cuentas: int, entidad_nombre: str) -> str:
    """Crea un README descriptivo para el ZIP exportado."""
    hoy = date.today().strftime("%d/%m/%Y")
    contenido = f"""
    CUENTAS PARTES - PAQUETE DE LIQUIDACIÓN (30 MESES)
    ==================================================

    Entidad: {entidad_nombre} (NIT: {entidad_nit})
    Fecha de generación: {hoy}

    Contenido:
    - CONSOLIDADO_GLOBAL.pdf (resumen ejecutivo y totales)
    - PDFs individuales por pensionado (30 cuentas por persona)

    Alcance metodológico:
    - Sistema de 30 cuentas independientes (mes vencido)
    - Capital fijo por cuenta, sin capitalización entre meses
    - Intereses por DTF mensual específica
    - Primas en junio y diciembre según número de mesadas

    Resumen del paquete:
    - Pensionados incluidos: {total_pensionados}
    - Total de cuentas independientes: {total_cuentas}

    Observaciones:
    - Septiembre no se factura en este corte; ventana: últimos 30 meses hasta agosto de 2025.
    - Este paquete es de uso interno y soporte de cobro persuasivo.
    """
    return "\n".join(line.rstrip() for line in contenido.splitlines()).strip() + "\n"

def generar_consolidado_global_texto(entidad_nit, entidad_nombre, todas_las_cuentas, total_capital, total_intereses, fecha_corte):
    """Genera el contenido del consolidado global en formato texto"""
    content = f"""
╔══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╗
║                                    CONSOLIDADO GLOBAL - LIQUIDACIÓN 30 CUENTAS INDEPENDIENTES                    ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                                                                                                  ║
║  🏛️  ENTIDAD: {entidad_nombre:<70}                                        ║
║  🆔  NIT: {entidad_nit:<77}                                        ║
║  📅  FECHA LIQUIDACIÓN: {fecha_corte.strftime('%d/%m/%Y'):<65}                                        ║
║  📊  PERÍODO: SEPTIEMBRE 2022 - FEBRERO 2025 (30 MESES EXACTOS)                                               ║
║                                                                                                                  ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                           RESUMEN EJECUTIVO                                                      ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                                                                                                  ║
║  👥  TOTAL PENSIONADOS: {len(todas_las_cuentas):<2}                                                                         ║
║  📋  TOTAL CUENTAS INDEPENDIENTES: {sum(len(p['cuentas']) for p in todas_las_cuentas):<3}                                                         ║
║  💰  CAPITAL TOTAL: ${float(total_capital):>15,.2f}                                                         ║
║  📈  INTERESES TOTALES: ${float(total_intereses):>15,.2f}                                                         ║
║  💯  GRAN TOTAL A COBRAR: ${float(total_capital + total_intereses):>15,.2f}                                                      ║
║                                                                                                                  ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                        RANKING POR PENSIONADO                                                    ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣

"""
    
    # Agregar ranking de pensionados
    pensionados_ordenados = sorted(todas_las_cuentas, key=lambda x: float(x['total_pensionado']), reverse=True)
    
    for i, p in enumerate(pensionados_ordenados, 1):
        porcentaje_del_total = (float(p['total_pensionado']) / float(total_capital + total_intereses)) * 100
        content += f"  {i:2d}. {p['pensionado']['nombre']:<35} ${float(p['total_pensionado']):>15,.2f} ({porcentaje_del_total:5.1f}%)\n"
    
    content += f"""

╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                          METODOLOGÍA APLICADA                                                    ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                                                                                                  ║
║  📋  SISTEMA: 30 Cuentas de Cobro Independientes por Pensionado                                                ║
║  💰  CAPITAL: Fijo por cuenta (sin capitalización entre cuentas)                                               ║
║  📈  INTERESES: DTF mensual específica aplicada sobre capital fijo                                             ║
║  🏛️  CONCEPTO: Mes vencido - Cuentas históricas desde su propio mes                                            ║
║  🎁  PRIMAS: Incluidas en diciembre y junio según número de mesadas                                            ║
║  ⚖️  MARCO LEGAL: Sistema Anti-Prescripción según Ley 1066 de 2006                                             ║
║                                                                                                                  ║
╚══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╝

NOTA IMPORTANTE: Este consolidado representa la suma de {sum(len(p['cuentas']) for p in todas_las_cuentas)} cuentas de cobro independientes.
Cada pensionado tiene exactamente 30 cuentas mensuales según prescripción vigente.

Generado el {fecha_corte.strftime('%d/%m/%Y')} - Sistema Cuotas Partes v2.0
"""
    
    return content

def generar_resumen_pensionado_texto(pensionado_data):
    """Genera el resumen individual de un pensionado"""
    p = pensionado_data['pensionado']
    
    content = f"""
╔══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╗
║                                     RESUMEN INDIVIDUAL - 30 CUENTAS INDEPENDIENTES                              ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                                                                                                  ║
║  👤  PENSIONADO: {p['pensionado']['nombre']:<70}                                           ║
║  🆔  CÉDULA: {p['pensionado']['cedula']:<77}                                           ║
║  📊  PORCENTAJE CUOTA: {float(p['pensionado']['porcentaje_cuota'])*100:>6.2f}%                                                                ║
║  🎁  NÚMERO MESADAS: {p['pensionado']['mesadas']:<2}                                                                              ║
║                                                                                                                  ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                           TOTALES PENSIONADO                                                     ║
╠══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╣
║                                                                                                                  ║
║  📋  TOTAL CUENTAS: {len(pensionado_data['cuentas']):<2}                                                                            ║
║  💰  CAPITAL TOTAL: ${float(pensionado_data['total_capital']):>15,.2f}                                                         ║
║  📈  INTERESES TOTALES: ${float(pensionado_data['total_intereses']):>15,.2f}                                                         ║
║  💯  GRAN TOTAL: ${float(pensionado_data['total_pensionado']):>15,.2f}                                                            ║
║                                                                                                                  ║
╚══════════════════════════════════════════════════════════════════════════════════════════════════════════════════╝

DETALLE DE LAS 30 CUENTAS:
════════════════════════════════════════════════════════════════════════════════════════════════════════════════════

 #  │   Mes/Año   │    Capital Base    │      Prima       │   Capital Total   │    Intereses     │   Total Cuenta   │ Estado
────┼─────────────┼────────────────────┼──────────────────┼───────────────────┼──────────────────┼──────────────────┼─────────
"""
    
    for cuenta in pensionado_data['cuentas']:
        estado_emoji = "🎁" if cuenta['estado'] == '🎁 PRIMA' else "📈"
        prima_str = f"${float(cuenta['prima']):>13,.2f}" if cuenta['prima'] > 0 else f"{'':>16}"
        
        content += f"{cuenta['consecutivo']:2d}  │ {cuenta['mes']:02d}/{cuenta['año']} │ ${float(cuenta['capital_base']):>14,.2f} │ {prima_str} │ ${float(cuenta['capital_total']):>13,.2f} │ ${float(cuenta['intereses']):>12,.2f} │ ${float(cuenta['total_cuenta']):>12,.2f} │ {estado_emoji}\n"
    
    content += f"""
────┴─────────────┴────────────────────┴──────────────────┴───────────────────┴──────────────────┴──────────────────┴─────────

METODOLOGÍA:
• Cada cuenta es independiente con capital fijo (sin capitalización)
• Intereses calculados desde el mes de la cuenta hasta agosto 2025
• DTF efectiva anual específica para cada mes
• Primas incluidas según número de mesadas (12/13/14)

# (secciones duplicadas eliminadas)
• Los intereses se calculan con la DTF del mes de la cuenta (mes vencido)
• Este sistema evita la prescripción de cuotas partes (Ley 1066/2006)

🔧 SOPORTE TÉCNICO:
───────────────────────────────────────────────────────────────────────────────────────────────────────────────────

Sistema: Cuotas Partes v2.0
Metodología: 30 Cuentas Independientes
Fecha generación: {date.today().strftime('%d/%m/%Y')}

Para consultas técnicas sobre la liquidación, contactar al administrador del sistema.
"""
    
    return content


def generar_zip_masivo_completo(entidad_nit: str, entidad_nombre: str, todas_las_cuentas: list, fecha_corte: date, corregir_existentes: bool = False, duplicate_policy: str | None = None) -> bytes:
    """
    Crea un ZIP en memoria con la estructura completa:
    - README.txt
    - CONSOLIDADO_GLOBAL.pdf
    - Carpeta por pensionado:
        - Carpeta por año:
            - PDF de cuenta de cobro individual.
    """
    import io
    import zipfile
    import os
    from datetime import datetime
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    
    # Importar la función de generación de PDF individual y configurar política de consecutivo
    import generar_pdf_oficial as gpo
    from generar_pdf_oficial import generar_pdf_para_pensionado
    try:
        # True: si existe una cuenta previa para el mismo periodo y pensionado, se actualiza (misma consecutivo)
        # False: crea una nueva entrada con nuevo consecutivo
        gpo.CONSEC_CORRECCION = bool(corregir_existentes)
        gpo.CONSEC_OVERRIDE = None
    except Exception:
        pass

    # Política explícita (si llega): puede ser "Reliquidar (conservar consecutivo)", "Usar existente (no generar)", "Crear nueva versión (nuevo consecutivo)"
    if duplicate_policy:
        if duplicate_policy.startswith("Reliquidar"):
            gpo.CONSEC_CORRECCION = True
        elif duplicate_policy.startswith("Crear nueva"):
            gpo.CONSEC_CORRECCION = False
        # "Usar existente" se maneja por fuera, evitando generar cuando hay archivo

    def _sanitize(name: str) -> str:
        import re
        name = name.strip().replace(' ', '_')
        name = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9_-]", '', name)
        return name

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        top_dir = f"{_sanitize(entidad_nombre)}_{entidad_nit}"
        # 1. README
        readme_content = generar_readme_texto(
            entidad_nit,
            len(todas_las_cuentas),
            sum(len(p.get('cuentas', [])) for p in todas_las_cuentas),
            entidad_nombre,
        )
        zf.writestr(f"{top_dir}/README.txt", readme_content)

        # 1.1 Asegurar que 'todas_las_cuentas' tenga totales calculados si viene en estructura mínima
        try:
            from decimal import Decimal
            from datetime import date as _date
            from dateutil.relativedelta import relativedelta
            # Misma lógica usada en la UI
            from mostrar_liquidacion_36 import ajustar_base_por_ipc, calcular_interes_mensual_unico
            from scripts.liquidacion_36_cuentas_corregida import tiene_prima_mes

            fecha_limite = _date(fecha_corte.year, fecha_corte.month, 1)
            for p in todas_las_cuentas:
                cuentas = p.get('cuentas', [])
                # Si ya tienen capital_total, asumimos que están completas
                if cuentas and isinstance(cuentas[0], dict) and 'capital_total' in cuentas[0]:
                    # Asegurar totales agregados en el nivel del pensionado
                    p.setdefault('total_capital', sum(c.get('capital_total', 0) for c in cuentas))
                    p.setdefault('total_intereses', sum(c.get('intereses', 0) for c in cuentas))
                    p.setdefault('total_pensionado', p['total_capital'] + p['total_intereses'])
                    continue

                # Calcular totales desde estructura mínima (solo año/mes)
                info = p.get('pensionado', {})
                try:
                    base = float(info.get('base_calculo_cuota', 0.0) or 0.0)
                except Exception:
                    base = 0.0
                try:
                    porcentaje = float(info.get('porcentaje_cuota', 0.0) or 0.0)
                except Exception:
                    porcentaje = 0.0
                try:
                    mesadas = int(info.get('mesadas', 12) or 12)
                except Exception:
                    mesadas = 12

                total_capital = Decimal('0')
                total_intereses = Decimal('0')
                total_total = Decimal('0')
                consecutivo = 1
                for cta in cuentas:
                    try:
                        año = int(cta['año']); mes = int(cta['mes'])
                    except Exception:
                        continue
                    fecha_cuenta = _date(año, mes, 1)
                    base_ajustada_año = ajustar_base_por_ipc(base, año)
                    capital_base = Decimal(str(base_ajustada_año)) * Decimal(str(porcentaje))
                    prima = capital_base if tiene_prima_mes(mesadas, mes) else Decimal('0')
                    capital_total = capital_base + prima

                    interes_acumulado = Decimal('0')
                    fa = fecha_cuenta
                    while fa <= fecha_limite:
                        im = calcular_interes_mensual_unico(float(capital_total), fa, fecha_corte)
                        interes_acumulado += Decimal(str(im))
                        fa = fa + relativedelta(months=1)

                    total_cuenta = capital_total + interes_acumulado

                    # Enriquecer la cuenta mínima con campos completos
                    cta['consecutivo'] = consecutivo
                    cta['capital_base'] = capital_base
                    cta['prima'] = prima
                    cta['capital_total'] = capital_total
                    cta['intereses'] = interes_acumulado
                    cta['total_cuenta'] = total_cuenta
                    cta['estado'] = '🎁 PRIMA' if prima > 0 else '📈 Regular'

                    total_capital += capital_total
                    total_intereses += interes_acumulado
                    total_total += total_cuenta
                    consecutivo += 1

                p['total_capital'] = total_capital
                p['total_intereses'] = total_intereses
                p['total_pensionado'] = total_total
        except Exception:
            # Si algo falla, seguimos sin bloquear la exportación; el consolidado puede quedar en cero para esos casos
            pass

        # Nota: Eliminamos la numeración visual por carpeta. A partir de ahora,
        # el número visible en cada PDF será el consecutivo global asignado y registrado en BD.

        # 2. PDF Consolidado: usar la MISMA función del botón "PDF Consolidado"
        #    Preferimos reutilizar el PDF ya generado en la sesión; si no existe, lo generamos aquí.
        temp_dir = "temp_pdf_generation"
        os.makedirs(temp_dir, exist_ok=True)
        error_log_path = os.path.join(temp_dir, "error_log.txt")

        pdf_consolidado_bytes = None
        pdf_consolidado_name = None
        try:
            # Generar SIEMPRE el consolidado fresco por entidad para asegurar numeración visible y contenido correctos
            pdf_consolidado_bytes, pdf_consolidado_name, cons_consolidado = generar_pdf_consolidado_en_memoria(
                entidad_nit=entidad_nit,
                entidad_nombre=entidad_nombre,
                todas_las_cuentas=todas_las_cuentas,
                fecha_corte=fecha_corte,
            )

            if pdf_consolidado_bytes and pdf_consolidado_name:
                # Agregar el consolidado como archivo independiente
                zf.writestr(os.path.join(top_dir, pdf_consolidado_name), pdf_consolidado_bytes)
        except Exception as e:
            with open(error_log_path, "a", encoding="utf-8") as f:
                f.write(f"Error generando/anexando PDF consolidado de la entidad {entidad_nombre} (NIT: {entidad_nit}): {e}\n")

        # 2.1 Generar Carta y añadir PDF combinado (Carta + Consolidado)
        try:
            template_docx = os.path.join(os.path.dirname(__file__), "Modelo cuenta de cobro cuotas partes pensionales (1).docx")
            if os.path.exists(template_docx) and pdf_consolidado_bytes:
                carta_docx_bytes, _carta_name = generar_carta_docx_asunto_en_memoria(
                    template_path=template_docx,
                    entidad_nombre=entidad_nombre,
                    fecha_corte=fecha_corte,
                    todas_las_cuentas=todas_las_cuentas,
                )
                carta_pdf_bytes = convert_docx_bytes_to_pdf_bytes(carta_docx_bytes)
                combinado_bytes = unir_pdfs_en_memoria(carta_pdf_bytes, pdf_consolidado_bytes)
                combinado_name = f"CARTA_Y_CONSOLIDADO_{entidad_nit}_{fecha_corte.strftime('%Y%m%d')}.pdf"
                zf.writestr(os.path.join(top_dir, combinado_name), combinado_bytes)
        except Exception as e:
            with open(error_log_path, "a", encoding="utf-8") as f:
                f.write(f"Error generando/anexando Carta+Consolidado: {e}\n")

        # 3. Generar PDFs individuales y organizarlos en carpetas

        # Reutilidad: construir nombre de carpeta del pensionado como lo hace el generador
        import re
        def _carpeta_pensionado_from_nombre_id(nombre: str, identificacion: str) -> str:
            nombre_completo = str(nombre or '')
            if ',' in nombre_completo:
                bloque_apellidos = nombre_completo.split(',')[0].strip()
            else:
                bloque_apellidos = nombre_completo.strip()
            partes = [p for p in bloque_apellidos.split() if p]
            ap1 = partes[0] if partes else ''
            ap2 = partes[1] if len(partes) > 1 else ''
            ap1 = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9_-]", '', ap1).strip()
            ap2 = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9_-]", '', ap2).strip()
            if ap1 and ap2:
                return f"{ap1.title()}_{ap2.title()}_{identificacion}"
            elif ap1:
                return f"{ap1.title()}_{identificacion}"
            else:
                return str(identificacion)

        # Base global de PDFs ya generados (fuera de temp)
        base_dir_global = os.path.join(os.path.dirname(__file__), 'reportes_liquidacion')

        # Sesión a BD para buscar existentes cuando la política sea "Usar existente"
        from app.db import get_session as _get_session
        from app.models import CuentaCobro as _CuentaCobro
        from sqlalchemy import select as _select
        from datetime import date as _date

        for pensionado_data in todas_las_cuentas:
            pensionado_info = pensionado_data['pensionado']
            
            # Reconstruir la tupla de pensionado que espera `generar_pdf_para_pensionado`
            # (identificacion, nombre, numero_mesadas, fecha_ingreso_nomina, empresa, base_calculo, porcentaje, nit)
            # Los índices pueden variar, es crucial ajustarlos al `pensionado_data` real.
            base_calc = pensionado_info.get('base_calculo_cuota', 0.0)
            try:
                base_calc = float(base_calc)
            except Exception:
                base_calc = 0.0
            porcentaje = pensionado_info.get('porcentaje_cuota', 0.0)
            try:
                porcentaje = float(porcentaje)
            except Exception:
                pass
            pensionado_tuple = (
                str(pensionado_info['cedula']),
                pensionado_info['nombre'],
                int(pensionado_info.get('mesadas', 12)),
                None,  # fecha_ingreso_nomina (no requerida por el cálculo)
                entidad_nombre,
                base_calc,
                porcentaje,
                str(entidad_nit),
            )

            for cuenta in pensionado_data['cuentas']:
                año = cuenta['año']
                mes = cuenta['mes']
                
                # Generar el PDF para una sola cuenta (un mes)
                # La función `generar_pdf_para_pensionado` crea el archivo en disco.
                try:
                    # Si la política es "Usar existente (no generar)", intentar localizar PDF previo
                    if duplicate_policy and duplicate_policy.startswith("Usar existente"):
                        periodo_inicio = _date(int(año), int(mes), 1)
                        periodo_fin = _date(fecha_corte.year, fecha_corte.month, 1)
                        with _get_session() as _s:
                            reg = _s.execute(
                                _select(_CuentaCobro).where(
                                    _CuentaCobro.nit_entidad == str(entidad_nit),
                                    _CuentaCobro.pensionado_identificacion == str(pensionado_info['cedula']),
                                    _CuentaCobro.periodo_inicio == periodo_inicio,
                                    _CuentaCobro.periodo_fin == periodo_fin,
                                ).order_by(_CuentaCobro.fecha_creacion.desc()).limit(1)
                            ).scalars().first()
                        if reg and reg.archivo_pdf:
                            carpeta_pensionado = _carpeta_pensionado_from_nombre_id(pensionado_info['nombre'], str(pensionado_info['cedula']))
                            posible_path = os.path.join(base_dir_global, carpeta_pensionado, reg.archivo_pdf)
                            if os.path.exists(posible_path):
                                # Añadir el archivo existente al ZIP como si fuera generado
                                base_name = os.path.basename(posible_path)
                                pensioner_folder_name = carpeta_pensionado
                                zip_path = os.path.join(top_dir, pensioner_folder_name, str(año), base_name)
                                zf.write(posible_path, arcname=zip_path)
                                continue  # No generar de nuevo

                    pdf_path = generar_pdf_para_pensionado(
                        pensionado=pensionado_tuple,
                        periodo='custom',
                        año_inicio=año,
                        mes_inicio=mes,
                        solo_mes=True, # ¡Importante para generar solo 1 cuenta!
                        output_dir=temp_dir, # Guardar en una carpeta temporal
                        display_consecutivo=cons_consolidado,
                        titulo_override='LIQUIDACION',
                        persistir_en_bd=False,
                    )
                    
                    if pdf_path and os.path.exists(pdf_path):
                        # Construir la ruta deseada dentro del ZIP
                        # Ej: Suarez_Mootoo_15240013/2023/15240013_Enero_2023.pdf
                        base_name = os.path.basename(pdf_path)
                        
                        # Extraer nombre de carpeta del pensionado desde la ruta generada
                        # La ruta es como temp_pdf_generation/Pensionado_ID/archivo.pdf
                        pensioner_folder_name = os.path.basename(os.path.dirname(pdf_path))
                        zip_path = os.path.join(top_dir, pensioner_folder_name, str(año), base_name)
                        
                        # Añadir el archivo al ZIP
                        zf.write(pdf_path, arcname=zip_path)
                        
                        # Opcional: eliminar el archivo temporal para no ocupar espacio
                        os.remove(pdf_path)

                except Exception as e:
                    # Registrar el error para no detener todo el proceso y continuar
                    error_log_path = os.path.join(temp_dir, "error_log.txt")
                    with open(error_log_path, "a", encoding="utf-8") as f:
                        f.write(f"Error generando PDF para {pensionado_info['cedula']} (Mes: {mes}/{año}): {e}\n")

        # Incluir el log de errores (si existe) dentro del ZIP para diagnóstico
        error_log_path = os.path.join(temp_dir, "error_log.txt")
        if os.path.exists(error_log_path):
            with open(error_log_path, "rb") as f:
                zf.writestr(f"{top_dir}/error_log.txt", f.read())

    # Limpiar la carpeta temporal
    import shutil
    shutil.rmtree(temp_dir, ignore_errors=True)
    
    return zip_buffer.getvalue()


# Construye lista de (año, mes) para los últimos MESES_PRESCRIPCION meses hasta fecha_corte (inclusive)
def _periodos_ultimos_meses(fecha_corte: date, meses: int) -> list[tuple[int,int]]:
    from dateutil.relativedelta import relativedelta
    inicio = date(fecha_corte.year, fecha_corte.month, 1) - relativedelta(months=meses-1)
    actual = inicio
    periodos = []
    while actual <= date(fecha_corte.year, fecha_corte.month, 1):
        periodos.append((actual.year, actual.month))
        actual = actual + relativedelta(months=1)
    return periodos


# Construye la estructura mínima esperada por generar_zip_masivo_completo para una entidad
def _construir_todas_cuentas_min(session, entidad_nit: str, fecha_corte: date) -> list:
    from sqlalchemy import text
    from app.settings import MESES_PRESCRIPCION
    periodos = _periodos_ultimos_meses(fecha_corte, MESES_PRESCRIPCION)
    rows = session.execute(text(
        """
        SELECT identificacion, nombre, numero_mesadas, base_calculo_cuota_parte, porcentaje_cuota_parte
        FROM pensionado WHERE nit_entidad = :nit ORDER BY nombre
        """
    ), {"nit": entidad_nit}).fetchall()
    resultado = []
    for r in rows:
        resultado.append({
            'pensionado': {
                'cedula': r[0],
                'nombre': r[1],
                'mesadas': int(r[2] or 12),
                'base_calculo_cuota': float(r[3] or 0),
                'porcentaje_cuota': float(r[4] or 0),
            },
            'cuentas': [ {'año': a, 'mes': m} for a, m in periodos ]
        })
    return resultado

# Utilidad para evitar desbordes de texto en tablas
def _shorten(text, max_len=28):
    try:
        s = str(text)
        return (s[: max_len - 1] + "…") if len(s) > max_len else s
    except Exception:
        return text

# --- Dashboard ---
if menu == "🏠 Dashboard":
    st.title("Generador de Cuentas de Cobro (Cuotas Partes)")
    st.subheader("Resumen global y por entidad")

    from app.db import get_session
    from sqlalchemy import text as _text
    import pandas as pd
    from datetime import datetime

    # Filtros
    with st.expander("Filtros"):
        c1, c2, c3 = st.columns([1, 1, 2])
        with c1:
            usar_rango = st.checkbox("Filtrar por período", value=False, help="Filtra por periodo_inicio y periodo_fin de las cuentas emitidas")
        with c2:
            hoy = date.today()
            if usar_rango:
                desde = st.date_input("Desde (inicio)", value=date(hoy.year, 1, 1), key="dash_desde")
                hasta = st.date_input("Hasta (fin)", value=hoy, key="dash_hasta")
            else:
                desde = None
                hasta = None
        with c3:
            # Cargar entidades para filtro
            try:
                session_tmp = get_session()
                ents = session_tmp.execute(_text("SELECT nit, nombre FROM entidad ORDER BY nombre")).fetchall()
                session_tmp.close()
            except Exception:
                ents = []
            etiquetas = [f"{e[1]} ({e[0]})" for e in ents]
            mapa_etq_a_nit = {f"{e[1]} ({e[0]})": str(e[0]) for e in ents}
            seleccion = st.multiselect("Entidades", options=etiquetas)
            nits_sel = [mapa_etq_a_nit[x] for x in seleccion if x in mapa_etq_a_nit]

        # Filtro por estado (EMITIDA, CORREGIDA, ANULADA, etc.)
        c4, _ = st.columns([1, 3])
        with c4:
            try:
                session_tmp2 = get_session()
                estados_rows = session_tmp2.execute(_text("SELECT DISTINCT estado FROM cuenta_cobro WHERE estado IS NOT NULL ORDER BY estado")).fetchall()
                session_tmp2.close()
                estados = [r[0] for r in estados_rows if r and r[0]]
            except Exception:
                estados = []
            estados_sel = st.multiselect("Estado", options=estados, default=estados)

    # Consultas agregadas
    try:
        session = get_session()

        # Construir WHERE dinámico
        where = ["1=1"]
        params = {}
        if usar_rango and desde:
            where.append("cc.periodo_inicio >= :desde")
            params["desde"] = desde
        if usar_rango and hasta:
            where.append("cc.periodo_fin <= :hasta")
            params["hasta"] = date(hasta.year, hasta.month, 1)
        if nits_sel:
            # Construir lista de marcadores para IN (...) evitando expansión de tuplas en SQL text
            marcadores = []
            for i, nitv in enumerate(nits_sel):
                key = f"nit{i}"
                marcadores.append(f":{key}")
                params[key] = nitv
            where.append(f"cc.nit_entidad IN ({', '.join(marcadores)})")
        # Estados seleccionados
        if 'estados_sel' in locals() and estados_sel and len(estados_sel) != 0 and (len(estados_sel) !=  len(set(estados_sel)) or True):
            marc_e = []
            for j, ev in enumerate(estados_sel):
                key = f"est{j}"
                marc_e.append(f":{key}")
                params[key] = ev
            where.append(f"cc.estado IN ({', '.join(marc_e)})")
        where_sql = " AND ".join(where)

        # Globales
        q_global = f"""
            SELECT COUNT(*) as cuentas,
                   COALESCE(SUM(cc.total_capital), 0) as capital,
                   COALESCE(SUM(cc.total_intereses), 0) as intereses,
                   COALESCE(SUM(cc.total_liquidacion), 0) as total
            FROM cuenta_cobro cc
            WHERE {where_sql}
        """
        g = session.execute(_text(q_global), params).fetchone()
        total_cuentas = int(g[0] or 0)
        total_capital = float(g[1] or 0)
        total_intereses = float(g[2] or 0)
        gran_total = float(g[3] or 0)

        # Por entidad
        q_ent = f"""
            SELECT cc.nit_entidad,
                   COALESCE(MAX(e.nombre), cc.nit_entidad) as nombre,
                   COUNT(*) as cuentas,
                   COALESCE(SUM(cc.total_capital),0) as capital,
                   COALESCE(SUM(cc.total_intereses),0) as intereses,
                   COALESCE(SUM(cc.total_liquidacion),0) as total,
                   MAX(cc.fecha_actualizacion) as actualizado
            FROM cuenta_cobro cc
            LEFT JOIN entidad e ON e.nit = cc.nit_entidad
            WHERE {where_sql}
            GROUP BY cc.nit_entidad
            ORDER BY total DESC
        """
        rows = session.execute(_text(q_ent), params).fetchall()

        # Métricas principales
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("🧾 Cuentas emitidas", f"{total_cuentas:,}")
        with col2:
            st.metric("💰 Capital total", f"${total_capital:,.0f}")
        with col3:
            st.metric("📈 Intereses", f"${total_intereses:,.0f}")
        with col4:
            st.metric("💯 Gran total", f"${gran_total:,.0f}")

        # KPIs del mes seleccionado (usa 'hasta' si hay rango; de lo contrario, mes actual)
        try:
            mes_ref = hasta if (usar_rango and hasta) else date.today()
            mes_key = date(mes_ref.year, mes_ref.month, 1)
            q_mes = f"""
                SELECT COUNT(*) as cuentas,
                       COALESCE(SUM(cc.total_capital),0) as capital,
                       COALESCE(SUM(cc.total_intereses),0) as intereses,
                       COALESCE(SUM(cc.total_liquidacion),0) as total
                FROM cuenta_cobro cc
                WHERE {where_sql} AND cc.periodo_fin = :meskey
            """
            params_mes = dict(params)
            params_mes["meskey"] = mes_key
            gm = session.execute(_text(q_mes), params_mes).fetchone()
            if gm:
                st.caption(f"KPIs del mes {mes_key.strftime('%Y-%m')}")
                m1, m2, m3, m4 = st.columns(4)
                with m1:
                    st.metric("Cuentas del mes", f"{int(gm[0] or 0):,}")
                with m2:
                    st.metric("Capital del mes", f"${float(gm[1] or 0):,.0f}")
                with m3:
                    st.metric("Intereses del mes", f"${float(gm[2] or 0):,.0f}")
                with m4:
                    st.metric("Total del mes", f"${float(gm[3] or 0):,.0f}")
        except Exception:
            pass

        st.markdown("---")
        st.subheader("📊 Resumen por entidad")
        if rows:
            df = pd.DataFrame(rows, columns=["NIT", "Entidad", "Cuentas", "Capital", "Intereses", "Total", "Actualizado"])
            # Formateo de números para visualización, mantener datos originales para gráficos
            df_show = df.copy()
            for col in ["Capital", "Intereses", "Total"]:
                df_show[col] = df_show[col].apply(lambda x: f"${float(x):,.0f}")
            st.dataframe(df_show, use_container_width=True, height=420)

            # Exportaciones
            cexp1, cexp2 = st.columns(2)
            # CSV
            csv_bytes = df.to_csv(index=False).encode("utf-8-sig")
            with cexp1:
                st.download_button(
                    label="⬇️ Descargar CSV",
                    data=csv_bytes,
                    file_name="resumen_por_entidad.csv",
                    mime="text/csv",
                    key="dl_csv_resumen_entidad"
                )
            # Excel (fallback a CSV si falla)
            try:
                import io
                from pandas import ExcelWriter
                excel_buffer = io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine="xlsxwriter") as writer:
                    df.to_excel(writer, index=False, sheet_name="Resumen")
                excel_data = excel_buffer.getvalue()
                with cexp2:
                    st.download_button(
                        label="⬇️ Descargar Excel",
                        data=excel_data,
                        file_name="resumen_por_entidad.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="dl_xlsx_resumen_entidad"
                    )
            except Exception:
                # Si no hay motor Excel, dejamos solo CSV
                pass

            # Top 10 gráfico
            with st.expander("Top 10 por total"):
                try:
                    df_top = df.sort_values("Total", ascending=False).head(10)[["Entidad", "Total"]]
                    df_top = df_top.set_index("Entidad")
                    st.bar_chart(df_top)
                except Exception:
                    pass
        else:
            st.info("No hay cuentas emitidas que coincidan con los filtros.")

        # Últimas cuentas generadas
        st.markdown("---")
        st.subheader("🕒 Últimas cuentas generadas")
        try:
            q_last = f"""
                SELECT cc.consecutivo, COALESCE(e.nombre, cc.nit_entidad) as entidad, cc.pensionado_nombre,
                       cc.pensionado_identificacion, cc.periodo_inicio, cc.periodo_fin,
                       cc.total_liquidacion, cc.fecha_actualizacion
                FROM cuenta_cobro cc
                LEFT JOIN entidad e ON e.nit = cc.nit_entidad
                WHERE {where_sql}
                ORDER BY cc.fecha_actualizacion DESC
                LIMIT 50
            """
            last_rows = session.execute(_text(q_last), params).fetchall()
            if last_rows:
                df_last = pd.DataFrame(last_rows, columns=[
                    "Consecutivo", "Entidad", "Pensionado", "Cédula", "Inicio", "Fin", "Total", "Actualizado"
                ])
                st.dataframe(df_last, use_container_width=True, height=360)
            else:
                st.caption("Sin registros recientes.")
        except Exception:
            pass

        session.close()
    except Exception as e:
        st.error(f"Error al obtener datos del Dashboard: {e}")
        st.info("Verifica la conexión a la base de datos y vuelve a intentarlo.")

# --- Módulo Pensionados ---
elif menu == "👤 Pensionados":
    st.title("Gestión de Pensionados")
    st.subheader("Registro / Edición de pensionados")
    
    try:
        from app.db import get_session
        from app.models import Pensionado
        import pandas as pd
        
        session = get_session()
        pensionados = session.query(Pensionado).all()
        session.close()
        
        st.markdown("---")
        # Buscador funcional
        filtro = st.text_input("🔍 Buscar pensionado por nombre, identificación o entidad")
        
        if pensionados:
            df = pd.DataFrame([
                {
                    "ID": p.pensionado_id,
                    "Identificación": p.identificacion,
                    "Nombre": p.nombre,
                    "Entidad": p.nit_entidad,
                    "Estado": p.estado_cartera,
                    "Fecha ingreso": p.fecha_ingreso_nomina,
                    "Capital pendiente": p.capital_pendiente,
                    "Intereses pendientes": p.intereses_pendientes
                }
                for p in pensionados
            ])
            
            if filtro:
                filtro_lower = filtro.lower()
                df = df[df.apply(lambda row: filtro_lower in str(row["Nombre"]).lower() or filtro_lower in str(row["Identificación"]).lower() or filtro_lower in str(row["Entidad"]).lower(), axis=1)]
            
            st.markdown('<div style="overflow-x:auto; max-width:100vw;">', unsafe_allow_html=True)
            st.dataframe(df, use_container_width=True, height=400)
            st.markdown('</div>', unsafe_allow_html=True)
            st.caption("Usa la barra inferior para deslizar horizontalmente y ver todas las columnas.")
        else:
            st.info("No hay pensionados registrados en la base de datos.")
            
    except Exception as e:
        st.error(f"Error al conectar con la base de datos: {e}")
        st.info("Verifica que MySQL esté ejecutándose y que la configuración sea correcta.")

# --- Módulo Liquidaciones ---
elif menu == "📑 Liquidaciones":
    st.title("Generar Liquidaciones Mensuales")
    st.subheader("Formulario de liquidación")
    
    try:
        from app.db import get_session
        from app.models import Pensionado, Entidad
        import pandas as pd

        session = get_session()
        entidades = session.query(Entidad).all()
        pensionados = session.query(Pensionado).all()
        session.close()

        entidad_sel = st.selectbox(
            "Selecciona entidad",
            entidades,
            format_func=lambda e: f"{e.nombre} ({e.nit})" if e else "",
            key="liq_entidad_sel"
        )
        pensionado_sel = st.selectbox(
            "Selecciona pensionado",
            pensionados,
            format_func=lambda p: f"{p.nombre} ({p.identificacion})" if p else "",
            key="liq_pensionado_sel"
        )
        periodo = st.date_input("Período a liquidar", value=date.today(), key="liq_periodo")
        valor_base = st.number_input("Valor base de cálculo", min_value=0.0, step=1000.0, key="liq_valor_base")
        tasa_dtf = st.number_input("Interés DTF (%)", min_value=0.0, step=0.01, key="liq_tasa_dtf")
        calcular = st.button("Liquidar cuenta", key="liq_btn_liquidar")

        st.markdown("---")

        if calcular and pensionado_sel and entidad_sel:
            # Simulación de cálculo real
            capital = valor_base
            interes = capital * (tasa_dtf / 100)
            total = capital + interes

            df = pd.DataFrame([
                {
                    "Pensionado": pensionado_sel.nombre,
                    "Identificación": pensionado_sel.identificacion,
                    "Entidad": entidad_sel.nombre,
                    "Periodo": periodo.strftime('%Y-%m'),
                    "Capital": f"{capital:,.2f}",
                    "Interés": f"{interes:,.2f}",
                    "Total": f"{total:,.2f}"
                }
            ])

            st.subheader(":mag: Vista previa de liquidación")
            st.dataframe(df, use_container_width=True)
            st.success(f"Liquidación generada: Capital={capital:,.2f}, Interés={interes:,.2f}, Total={total:,.2f}")
        elif calcular:
            st.error("Debes seleccionar entidad y pensionado.")
    except Exception as e:
        st.error(f"Error al conectar con la base de datos: {e}")
        st.info("Verifica que MySQL esté ejecutándose y que la configuración sea correcta.")

# --- Módulo Pagos ---
elif menu == "💰 Pagos":
    st.title("Registro de Pagos de Entidades")
    st.subheader("Formulario de registro de pagos")
    
    st.text_input("Seleccionar pensionado")
    st.text_input("Seleccionar entidad")
    st.text_input("Número de cuenta de cobro aplicada")
    st.date_input("Fecha de pago", value=date.today())
    st.number_input("Valor abonado", min_value=0.0, step=1000.0)
    st.text_area("Observaciones")
    st.button("Guardar pago")
    
    st.markdown("---")
    st.write("(Aquí se mostraría la tabla de pagos registrados)")

# --- Módulo Cobro Persuasivo ---
elif menu == "📤 Cobro Persuasivo":
    st.title("Cobro Persuasivo a Entidades Deudoras")
    st.subheader("Generar carta de cobro")
    
    st.text_input("Seleccionar entidad deudora")
    st.date_input("Período a cobrar", value=date.today())
    st.text_area("Información legal (Ley 1066/2006)")
    st.write("Valores y períodos a cobrar (simulado)")
    st.file_uploader("Adjuntar liquidaciones mensuales + consolidado", accept_multiple_files=True)
    st.button("Enviar / Descargar carta")

# --- Módulo Reportes y Seguimiento ---
elif menu == "📊 Reportes y Seguimiento":
    st.title("Reportes y Seguimiento")
    st.subheader("Tabla dinámica de seguimiento")
    
    st.text_input("Filtrar por entidad")
    st.text_input("Filtrar por pensionado")
    st.date_input("Filtrar por período", value=date.today())
    
    st.markdown("---")
    st.write("(Aquí se mostraría la tabla de reportes y seguimiento)")
    st.button("Exportar a PDF")
    st.button("Exportar a Excel")
    
    st.markdown("---")
    st.write("🔴 Cuentas próximas a prescribir | 🟡 Pagos parciales | 🟢 Pagos completos")

# --- Módulo Liquidaciones Masivas (30 Cuentas) ---
elif menu == "⚖️ Liquidaciones Masivas (30 Cuentas)":
    st.title("⚖️ Liquidaciones Masivas (30 Cuentas Independientes)")
    
    st.markdown("""
    ### 📋 Sistema Corregido de 30 Cuentas Independientes
    **🎯 Objetivo:** Generar 30 cuentas de cobro independientes por pensionado (mes vencido).
    
    **💡 Metodología Corregida:**
    - **📅 Período:** Últimos 30 meses hasta agosto de 2025 (septiembre no se toma por facturar)
    - **💰 Capital Fijo:** Cada cuenta mantiene su capital base sin capitalización
    - **📈 Interés Simple:** DTF mensual aplicada sobre capital fijo de cada cuenta
    - **🏛️ Mes Vencido:** Cuentas históricas generan intereses desde su propio mes
    - **🎁 Primas:** Incluidas en diciembre y junio según corresponda
    
    **⚖️ Marco Legal:** Sistema Anti-Prescripción - Cada mes es una cuenta independiente
    """)
    
    # Inicializar session_state
    if 'cuentas_generadas' not in st.session_state:
        st.session_state.cuentas_generadas = None
        st.session_state.entidad_actual = None
    
    try:
        from app.db import get_session
        from sqlalchemy import text
        import os
        
        # Importar las funciones del sistema corregido (desde el paquete scripts)
        from scripts.liquidacion_36_cuentas_corregida import (
            obtener_pensionados_entidad, 
            calcular_cuenta_mensual, 
            obtener_dtf_mes,
            tiene_prima_mes,
            calcular_interes_mensual,
            calcular_dias_mes
        )
        from mostrar_liquidacion_36 import ajustar_base_por_ipc
        # Usar la misma fórmula de intereses que el PDF de "solo mes" para alinear valores
        from mostrar_liquidacion_36 import calcular_interes_mensual_unico
        
        session = get_session()
        
        # Obtener entidades disponibles
        entidades = session.execute(
            text("""
                SELECT DISTINCT e.nit, e.nombre, COUNT(p.pensionado_id) as pensionados_activos
                FROM entidad e
                LEFT JOIN pensionado p ON e.nit = p.nit_entidad AND p.estado_cartera = 'ACTIVO'
                GROUP BY e.nit, e.nombre
                HAVING pensionados_activos > 0
                ORDER BY e.nombre
            """)
        ).fetchall()
        
        st.markdown("---")
        
        # Formulario de generación
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.subheader("🏢 Selección de Entidad")
            
            if entidades:
                entidad_options = [(e.nit, f"{e.nombre} (NIT: {e.nit}) - {e.pensionados_activos} pensionados") for e in entidades]
                entidad_selected = st.selectbox(
                    "Selecciona la entidad:",
                    options=entidad_options,
                    format_func=lambda x: x[1],
                    key="masivas_entidad"
                )
                entidad_nit = entidad_selected[0] if entidad_selected else None
                # Obtener nombre de la entidad seleccionada
                entidad_nombre = next((e.nombre for e in entidades if e.nit == entidad_nit), "ENTIDAD NO ENCONTRADA")
            else:
                st.error("No se encontraron entidades con pensionados activos")
                entidad_nit = None
                entidad_nombre = None
        
        with col2:
            st.subheader("📅 Configuración")
            fecha_corte = st.date_input(
                "Fecha de corte:",
                value=date(2025, 8, 31),
                help="Usado para calcular los últimos 30 meses (mes vencido).",
                key="masivas_fecha_corte"
            )
            
            st.info(f"""
            **📆 Período (dinámico):**
            - Últimos 30 meses hasta la fecha de corte seleccionada
            - Con corte agosto 2025: 01/03/2023 → 31/08/2025 (30 meses)
            - 💰 **Metodología:** Mes vencido
            """)
        
        st.markdown("---")
        
        # Botón de generación
        if st.button("🚀 Generar 30 Cuentas de Cobro", type="primary"):
            if entidad_nit:
                with st.spinner("⏳ Generando las 30 cuentas independientes..."):
                    try:
                        from decimal import Decimal, InvalidOperation
                        from datetime import date
                        from dateutil.relativedelta import relativedelta
                        
                        # Obtener pensionados de la entidad
                        pensionados = obtener_pensionados_entidad(session, entidad_nit)
                        
                        if not pensionados:
                            st.error(f"No se encontraron pensionados para la entidad {entidad_nit}")
                        else:
                            # Período dinámico: últimos 30 meses hasta fecha_corte
                            from dateutil.relativedelta import relativedelta
                            fecha_final = date(fecha_corte.year, fecha_corte.month, 1)
                            fecha_inicial = fecha_final - relativedelta(months=29)
                            
                            # Generar todas las cuentas
                            todas_las_cuentas = []
                            total_capital_entidad = Decimal('0')
                            total_intereses_entidad = Decimal('0')
                            
                            # Barra de progreso
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            pensionado_count = 0
                            total_pensionados = len(pensionados)
                            
                            for pensionado in pensionados:
                                pensionado_count += 1
                                status_text.text(f"Procesando pensionado {pensionado_count}/{total_pensionados}: {pensionado[1]}")
                                
                                # Generar 30 cuentas para este pensionado
                                fecha_cuenta = fecha_inicial
                                consecutivo = 1
                                total_pensionado = Decimal('0')
                                
                                cuentas_pensionado = []

                                # Base cálculo cuota original del pensionado (para mostrar en consolidado)
                                try:
                                    base_calculo_cuota_parte_origen = float(pensionado[8]) if pensionado[8] else 383628.0
                                except (ValueError, IndexError, TypeError):
                                    base_calculo_cuota_parte_origen = 383628.0
                                
                                while fecha_cuenta <= fecha_final:
                                    # Obtener datos del pensionado de forma segura
                                    try:
                                        # Usar porcentaje_cuota del campo correcto (índice 3)
                                        porcentaje_cuota = Decimal(str(pensionado[3])) if pensionado[3] else Decimal('0.15')
                                        # Usar numero_mesadas del campo correcto (índice 4) 
                                        numero_mesadas = int(pensionado[4]) if pensionado[4] else 12
                                        # Usar base_calculo_cuota_parte de la base de datos (índice 8)
                                        base_calculo_cuota_parte = Decimal(str(pensionado[8])) if pensionado[8] else Decimal('383628.0')
                                    except (ValueError, IndexError, InvalidOperation):
                                        # Valores por defecto en caso de error
                                        porcentaje_cuota = Decimal('0.15')
                                        numero_mesadas = 12
                                        base_calculo_cuota_parte = Decimal('922628.0')  # Valor correcto de la BD
                                    
                                    # Ajustar base por IPC (base 2025 → año de la cuenta) usando la misma función que el cálculo individual
                                    base_ajustada_año = ajustar_base_por_ipc(float(base_calculo_cuota_parte), fecha_cuenta.year)
                                    
                                    # Calcular capital fijo de la cuenta usando la base ajustada por IPC
                                    capital_base = Decimal(str(base_ajustada_año)) * porcentaje_cuota
                                    
                                    # Determinar si tiene prima
                                    tiene_prima = tiene_prima_mes(numero_mesadas, fecha_cuenta.month)
                                    prima = capital_base if tiene_prima else Decimal('0')
                                    capital_total = capital_base + prima
                                    
                                    # Calcular intereses acumulativos para ESTA cuenta específica
                                    fecha_actual = fecha_cuenta
                                    # Limitar intereses hasta el primer día del mes de corte
                                    fecha_limite = date(fecha_corte.year, fecha_corte.month, 1)
                                    interes_acumulado = Decimal('0')
                                    
                                    # Iterar mes a mes DESDE el mes de la cuenta hasta el mes de corte
                                    while fecha_actual <= fecha_limite:
                                        # Interés del mes usando la misma fórmula que el PDF individual (capital fijo)
                                        interes_mes = calcular_interes_mensual_unico(float(capital_total), fecha_actual, fecha_corte)
                                        interes_acumulado += Decimal(str(interes_mes))
                                        fecha_actual = fecha_actual + relativedelta(months=1)
                                    
                                    total_cuenta = capital_total + interes_acumulado
                                    
                                    cuenta = {
                                        'pensionado_id': pensionado[0],  # pensionado_id está en índice 0
                                        'pensionado_nombre': pensionado[2],  # nombre está en índice 2
                                        'cedula': pensionado[1],  # cedula está en índice 1
                                        'porcentaje': float(porcentaje_cuota),
                                        'consecutivo': consecutivo,
                                        'mes': fecha_cuenta.month,
                                        'año': fecha_cuenta.year,
                                        'fecha_cuenta': fecha_cuenta,
                                        'capital_base': capital_base,
                                        'prima': prima,
                                        'capital_total': capital_total,
                                        'intereses': interes_acumulado,
                                        'total_cuenta': total_cuenta,
                                        'estado': '🎁 PRIMA' if tiene_prima else '📈 Regular'
                                    }
                                    
                                    cuentas_pensionado.append(cuenta)
                                    total_pensionado += total_cuenta
                                    
                                    fecha_cuenta += relativedelta(months=1)
                                    consecutivo += 1
                                
                                # Agregar resumen del pensionado
                                resumen_pensionado = {
                                    'pensionado': {
                                        'cedula': pensionado[1],  # identificacion está en índice 1
                                        'nombre': pensionado[2],  # nombre está en índice 2
                                        'porcentaje_cuota': porcentaje_cuota,
                                        'mesadas': numero_mesadas,
                                        'pensionado_id': pensionado[0],  # pensionado_id está en índice 0
                                        'base_calculo_cuota': base_calculo_cuota_parte_origen,
                                        'ingreso_nomina': pensionado[5] if len(pensionado) > 5 else None,
                                        'resolucion': pensionado[9] if len(pensionado) > 9 else None
                                    },
                                    'cuentas': cuentas_pensionado,
                                    'total_capital': sum(c['capital_total'] for c in cuentas_pensionado),
                                    'total_intereses': sum(c['intereses'] for c in cuentas_pensionado),
                                    'total_pensionado': total_pensionado
                                }
                                
                                todas_las_cuentas.append(resumen_pensionado)
                                total_capital_entidad += resumen_pensionado['total_capital']
                                total_intereses_entidad += resumen_pensionado['total_intereses']
                                
                                # Actualizar progreso
                                progress_bar.progress(pensionado_count / total_pensionados)
                            
                            progress_bar.progress(1.0)
                            status_text.text("✅ Generación completada!")
                            
                            # Guardar resultados en session_state
                            st.session_state.cuentas_generadas = {
                                'todas_las_cuentas': todas_las_cuentas,
                                'total_capital_entidad': total_capital_entidad,
                                'total_intereses_entidad': total_intereses_entidad,
                                'total_cuentas_generadas': sum(len(p['cuentas']) for p in todas_las_cuentas)
                            }
                            st.session_state.entidad_actual = entidad_nit
                        
                    except Exception as e:
                        st.error(f"❌ Error generando las 30 cuentas de cobro: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())
            else:
                st.error("⚠️ Por favor selecciona una entidad")
        
        # Mostrar resultados si existen en session_state
        if st.session_state.cuentas_generadas and st.session_state.entidad_actual == entidad_nit:
            datos = st.session_state.cuentas_generadas
            todas_las_cuentas = datos['todas_las_cuentas']
            total_capital_entidad = datos['total_capital_entidad']
            total_intereses_entidad = datos['total_intereses_entidad']
            total_cuentas_generadas = datos['total_cuentas_generadas']
            
            st.success(f"✅ {total_cuentas_generadas} cuentas independientes generadas para {len(todas_las_cuentas)} pensionados!")
            # Acceso rápido al dashboard
            if st.button("📊 Ver Dashboard con resumen", key="btn_goto_dashboard_masivo"):
                st.session_state["menu_principal"] = "🏠 Dashboard"
                st.rerun()
            
            # Consolidado global
            st.markdown("---")
            st.subheader("🏆 CONSOLIDADO GLOBAL DE LA ENTIDAD")
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("👥 Pensionados", f"{len(todas_las_cuentas)}")
            with col2:
                st.metric("💰 Capital Total", f"${float(total_capital_entidad):,.2f}")
            with col3:
                st.metric("📈 Intereses Total", f"${float(total_intereses_entidad):,.2f}")
            with col4:
                st.metric("💯 GRAN TOTAL", f"${float(total_capital_entidad + total_intereses_entidad):,.2f}")
            
            # Resumen por pensionado
            st.markdown("---")
            st.subheader("📋 Resumen por Pensionado")
            
            import pandas as pd
            df_pensionados = pd.DataFrame([
                {
                    'Pensionado': p['pensionado']['nombre'],
                    'Cédula': p['pensionado']['cedula'],
                    'Base Cálculo': f"${float(p['pensionado'].get('base_calculo_cuota', 0.0)):,.2f}",
                    'Porcentaje': f"{float(p['pensionado']['porcentaje_cuota']) * 100:.2f}%",
                    'Cuentas': len(p['cuentas']),
                    'Capital': f"${float(p['total_capital']):,.2f}",
                    'Intereses': f"${float(p['total_intereses']):,.2f}",
                    'Total': f"${float(p['total_pensionado']):,.2f}"
                }
                for p in todas_las_cuentas
            ])
            
            st.dataframe(df_pensionados, use_container_width=True, height=400)
            
            # Mostrar detalle de un pensionado específico
            st.markdown("---")
            st.subheader("🔍 Detalle por Pensionado")
            
            # Selector de pensionado con key único para evitar reinicios
            pensionado_names = [(i, p['pensionado']['nombre']) for i, p in enumerate(todas_las_cuentas)]
            selected_pensionado_idx = st.selectbox(
                "Selecciona un pensionado para ver sus 30 cuentas:",
                options=[i for i, _ in pensionado_names],
                format_func=lambda x: pensionado_names[x][1],
                key="masivas_selector_pensionado"
            )
            
            if selected_pensionado_idx is not None:
                pensionado_detalle = todas_las_cuentas[selected_pensionado_idx]
                
                st.write(f"**👤 {pensionado_detalle['pensionado']['nombre']}** - Total: ${float(pensionado_detalle['total_pensionado']):,.2f}")
                
                # Mostrar las 30 cuentas
                df_cuentas = pd.DataFrame([
                    {
                        'Cuenta': f"{c['consecutivo']:02d}",
                        'Mes/Año': f"{c['mes']:02d}/{c['año']}",
                        'Capital Base': f"${float(c['capital_base']):,.2f}",
                        'Prima': f"${float(c['prima']):,.2f}" if c['prima'] > 0 else "-",
                        'Capital Total': f"${float(c['capital_total']):,.2f}",
                        'Intereses': f"${float(c['intereses']):,.2f}",
                        'Total Cuenta': f"${float(c['total_cuenta']):,.2f}",
                        'Estado': c['estado']
                    }
                    for c in pensionado_detalle['cuentas']
                ])
                
                st.dataframe(df_cuentas, use_container_width=True, height=600)
            
            # Sección de exportación
            st.markdown("---")
            st.subheader("📦 Exportar Documentos Oficiales")

            # Placeholders para evitar botones duplicados tras múltiples clics
            if 'pdf_consolidado_data' not in st.session_state:
                st.session_state['pdf_consolidado_data'] = None
            if 'pdf_consolidado_name' not in st.session_state:
                st.session_state['pdf_consolidado_name'] = None
            if 'zip_masivo_data' not in st.session_state:
                st.session_state['zip_masivo_data'] = None
            if 'zip_masivo_name' not in st.session_state:
                st.session_state['zip_masivo_name'] = None

            col1, col2 = st.columns(2)

            with col1:
                pdf_dl_placeholder = st.empty()
                zip_dl_placeholder = st.empty()

                # Render persistente en el mismo placeholder si ya existen datos en sesión
                if st.session_state.get('pdf_consolidado_data') and st.session_state.get('pdf_consolidado_name'):
                    pdf_dl_placeholder.download_button(
                        label="⬇️ Descargar PDF Consolidado",
                        data=st.session_state['pdf_consolidado_data'],
                        file_name=st.session_state['pdf_consolidado_name'],
                        mime="application/pdf",
                        key="download_pdf_consolidado_persist"
                    )
                if st.session_state.get('zip_masivo_data') and st.session_state.get('zip_masivo_name'):
                    zip_dl_placeholder.download_button(
                        label="⬇️ Descargar ZIP (carpetas por año)",
                        data=st.session_state['zip_masivo_data'],
                        file_name=st.session_state['zip_masivo_name'],
                        mime="application/zip",
                        key="download_zip_masivo_persist"
                    )

                if st.button("📄 PDF Consolidado", type="primary", use_container_width=True):
                    try:
                        pdf_data, pdf_name, _cons = generar_pdf_consolidado_en_memoria(
                            entidad_nit=entidad_nit,
                            entidad_nombre=next((e.nombre for e in entidades if e.nit == entidad_nit), str(entidad_nit)),
                            todas_las_cuentas=todas_las_cuentas,
                            fecha_corte=fecha_corte,
                        )
                        st.session_state['pdf_consolidado_data'] = pdf_data
                        st.session_state['pdf_consolidado_name'] = pdf_name

                        pdf_dl_placeholder.download_button(
                            label="⬇️ Descargar PDF Consolidado",
                            data=st.session_state['pdf_consolidado_data'],
                            file_name=st.session_state['pdf_consolidado_name'],
                            mime="application/pdf",
                            key="download_pdf_consolidado"
                        )
                        st.success(f"✅ PDF consolidado generado para {len(todas_las_cuentas)} pensionados")
                    except Exception as e:
                        st.error(f"Error generando PDF consolidado: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())
                
                # Política para duplicados
                dup_policy = st.selectbox(
                    "Si ya existe la cuenta del mismo periodo…",
                    [
                        "Preguntar antes",
                        "Reliquidar (conservar consecutivo)",
                        "Usar existente (no generar)",
                        "Crear nueva versión (nuevo consecutivo)",
                    ],
                    index=1,
                    key="dup_policy_entity"
                )

                # Pre-escaneo de duplicados en cuenta_cobro para advertir antes de generar
                try:
                    from sqlalchemy import text
                    from datetime import date as _date
                    periodo_fin_chk = _date(fecha_corte.year, fecha_corte.month, 1)
                    duplicados = 0
                    for p in todas_las_cuentas:
                        ced = str(p['pensionado']['cedula'])
                        for c in p.get('cuentas', []):
                            ini = _date(c['año'], c['mes'], 1)
                            cnt = session.execute(
                                text("SELECT COUNT(*) FROM cuenta_cobro WHERE nit_entidad=:n AND pensionado_identificacion=:id AND periodo_inicio=:ini AND periodo_fin=:fin"),
                                {"n": str(entidad_nit), "id": ced, "ini": ini, "fin": periodo_fin_chk}
                            ).scalar() or 0
                            if cnt:
                                duplicados += 1
                                break  # ya sabemos que este pensionado tiene al menos un duplicado
                    if duplicados:
                        msg = "Se encontraron cuentas existentes para {n} pensionado(s).".format(n=duplicados)
                        if dup_policy == "Preguntar antes":
                            st.warning(msg + " Selecciona una acción en la lista y vuelve a generar.")
                        else:
                            st.warning(msg + f" Se aplicará: {dup_policy}.")
                    else:
                        st.info("No se encontraron cuentas existentes para los periodos a generar.")
                except Exception:
                    pass

                # Consulta rápida de consecutivos (BD vs archivo)
                with st.expander("Consecutivos (BD vs archivo)"):
                    if st.button("🔄 Recalcular siguiente consecutivo", key="btn_recalc_cons_masivo"):
                        try:
                            from sqlalchemy import text as _tx
                            mx = session.execute(_tx("SELECT MAX(consecutivo) FROM cuenta_cobro")).scalar()
                        except Exception:
                            mx = None
                        try:
                            with open('ultimo_consecutivo.txt', 'r', encoding='utf-8') as f:
                                u = int(f.read().strip())
                        except Exception:
                            u = None
                        colc1, colc2 = st.columns(2)
                        with colc1:
                            st.metric("Siguiente BD", (mx or 0) + 1)
                            st.caption(f"MAX BD: {mx or 0}")
                        with colc2:
                            st.metric("Siguiente archivo", (u or 0) + 1)
                            st.caption(f"Archivo actual: {u if u is not None else 'N/A'}")

                # Botón para generar ZIP con la nueva estructura de carpetas por año
                if st.button("📁 Generar ZIP (carpetas por año)", key="zip_masivo_completo"):
                    try:
                        with st.spinner("⏳ Generando ZIP con estructura de carpetas por año..."):
                            # Obtener el nombre de la entidad para el nombre del archivo
                            entidad_nombre = next((e.nombre for e in entidades if e.nit == entidad_nit), "ENTIDAD")
                            # Si el usuario eligió 'Preguntar antes', frenamos para que elija otra opción
                            if dup_policy == "Preguntar antes":
                                st.warning("Selecciona cómo manejar las cuentas ya existentes y vuelve a pulsar el botón.")
                                st.stop()
                            
                            # Llamar a la nueva función que genera el ZIP con subcarpetas de año
                            zip_data = generar_zip_masivo_completo(
                                entidad_nit=entidad_nit,
                                entidad_nombre=entidad_nombre,
                                todas_las_cuentas=todas_las_cuentas,
                                fecha_corte=fecha_corte,
                                corregir_existentes=(dup_policy == "Reliquidar (conservar consecutivo)"),
                                duplicate_policy=dup_policy,
                            )

                            # Guardar en session_state y renderizar/actualizar botón único
                            st.session_state['zip_masivo_data'] = zip_data
                            st.session_state['zip_masivo_name'] = f"LIQUIDACION_MASIVA_{entidad_nit}_{fecha_corte.strftime('%Y%m%d')}.zip"

                            # Además: guardar automáticamente en la misma carpeta que "personalizados"
                            try:
                                import os
                                # Obtener nombre oficial de la entidad
                                entidad_row = session.execute(text("SELECT nombre FROM entidad WHERE nit = :nit"), {"nit": entidad_nit}).fetchone()
                                entidad_nombre_fs = (entidad_row[0] if entidad_row else str(entidad_nit))
                                prefijo = entidad_nombre_fs.strip().upper()[:3].replace(' ', '')
                                carpeta_entidad = f"{prefijo}_{entidad_nit}"
                                base_dir = os.path.join(os.path.dirname(__file__), 'reportes_liquidacion', carpeta_entidad)
                                os.makedirs(base_dir, exist_ok=True)
                                target_path = os.path.join(base_dir, st.session_state['zip_masivo_name'])
                                # Asegurar nombre único si existe
                                if os.path.exists(target_path):
                                    base, ext = os.path.splitext(target_path)
                                    suf = 2
                                    while os.path.exists(f"{base}_v{suf}{ext}"):
                                        suf += 1
                                    target_path = f"{base}_v{suf}{ext}"
                                with open(target_path, 'wb') as f:
                                    f.write(st.session_state['zip_masivo_data'])
                                st.info(f"También se guardó en: {target_path}")
                            except Exception as _e:
                                st.warning(f"No se pudo guardar automáticamente el ZIP en disco: {_e}")

                            zip_dl_placeholder.download_button(
                                label="⬇️ Descargar ZIP (carpetas por año)",
                                data=st.session_state['zip_masivo_data'],
                                file_name=st.session_state['zip_masivo_name'],
                                mime="application/zip",
                                key="download_zip_masivo"
                            )

                            st.success(f"✅ ZIP generado con {len(todas_las_cuentas)} pensionados y estructura de carpetas por año.")
                            
                    except Exception as e:
                        st.error(f"❌ Error generando el ZIP masivo: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())
            
            with col2:
                st.info("""
                **📋 Opciones de Exportación:**
                
                **📄 PDF Consolidado:**
                - Un solo documento con todos los pensionados
                - Tabla resumen con totales generales
                - Formato oficial cuenta de cobro
                - Ideal para presentaciones ejecutivas
                
                **📁 ZIP con PDFs Individuales:**
                - Un PDF por cada pensionado
                - Formato basado en plantilla Excel oficial
                - Numeración consecutiva automática
                - Estructura completa según PLANTILLA CXC 06
                """)

            st.markdown("---")
            # Carta a partir del modelo: generar DOCX y convertir a PDF
            st.subheader("✉️ Carta de Presentación (PDF)")
            try:
                template_docx = os.path.join(os.path.dirname(__file__), "Modelo cuenta de cobro cuotas partes pensionales (1).docx")
                if os.path.exists(template_docx):
                    if st.button("📝 Generar Carta (PDF)", key="btn_gen_carta_pdf"):
                        try:
                            carta_bytes, carta_name = generar_carta_docx_asunto_en_memoria(
                                template_path=template_docx,
                                entidad_nombre=next((e.nombre for e in entidades if e.nit == entidad_nit), str(entidad_nit)),
                                fecha_corte=fecha_corte,
                                todas_las_cuentas=todas_las_cuentas,
                            )
                            # Convertir DOCX -> PDF usando conversión centralizada robusta
                            try:
                                pdf_bytes = convert_docx_bytes_to_pdf_bytes(carta_bytes)
                                pdf_name = os.path.splitext(carta_name)[0] + ".pdf"
                                st.download_button(
                                    label="⬇️ Descargar Carta PDF",
                                    data=pdf_bytes,
                                    file_name=pdf_name,
                                    mime="application/pdf",
                                    key="dl_carta_pdf"
                                )
                                st.info("La carta se generó a partir del modelo y se convirtió a PDF.")
                            except ImportError:
                                st.error("No se encontró docx2pdf. Instálalo y asegúrate de tener Microsoft Word en Windows para convertir DOCX a PDF.")
                            except Exception as conv_err:
                                st.error(f"No fue posible convertir la carta a PDF: {conv_err}")
                            finally:
                                pass
                        except Exception as _e_docx:
                            st.error(f"No fue posible generar la carta DOCX: {_e_docx}")
                else:
                    st.caption("Modelo de carta no encontrado en la carpeta del proyecto.")
            except Exception:
                st.caption("No fue posible preparar la generación de la carta.")
            st.subheader("🚚 Generación Masiva por Todas las Entidades")
            
            # Botón para unir Carta + Consolidado sin modificar contenidos
            st.markdown("---")
            st.subheader("📎 Unir Carta + Consolidado (PDF)")
            st.caption("Genera un único PDF con la carta primero y el consolidado después, sin alterar nada.")
            if st.button("📌 Unir Carta y Consolidado", key="btn_merge_carta_consol"):
                try:
                    # 1. Generar carta PDF en memoria (reutilizamos la lógica de arriba)
                    template_docx = os.path.join(os.path.dirname(__file__), "Modelo cuenta de cobro cuotas partes pensionales (1).docx")
                    if not os.path.exists(template_docx):
                        raise RuntimeError("No se encontró el modelo de carta en el proyecto.")
                    carta_bytes_docx, carta_name = generar_carta_docx_asunto_en_memoria(
                        template_path=template_docx,
                        entidad_nombre=next((e.nombre for e in entidades if e.nit == entidad_nit), str(entidad_nit)),
                        fecha_corte=fecha_corte,
                        todas_las_cuentas=todas_las_cuentas,
                    )
                    # Conversión robusta centralizada
                    carta_pdf_bytes = convert_docx_bytes_to_pdf_bytes(carta_bytes_docx)

                    # 2. Generar PDF consolidado en memoria usando la función existente
                    cons_pdf_bytes, cons_pdf_name, _cons = generar_pdf_consolidado_en_memoria(
                        entidad_nit=entidad_nit,
                        entidad_nombre=next((e.nombre for e in entidades if e.nit == entidad_nit), str(entidad_nit)),
                        todas_las_cuentas=todas_las_cuentas,
                        fecha_corte=fecha_corte,
                    )

                    # 3. Unir ambos PDFs: carta primero, consolidado después
                    try:
                        merged_bytes = unir_pdfs_en_memoria(carta_pdf_bytes, cons_pdf_bytes)
                    except Exception as me:
                        raise RuntimeError(me)

                    merged_name = f"CARTA_Y_CONSOLIDADO_{entidad_nit}_{fecha_corte.strftime('%Y%m%d')}.pdf"
                    # Guardado automático en carpeta de la entidad (igual que el ZIP)
                    try:
                        import os
                        from sqlalchemy import text as _text
                        entidad_row = session.execute(_text("SELECT nombre FROM entidad WHERE nit = :nit"), {"nit": entidad_nit}).fetchone()
                        entidad_nombre_fs = (entidad_row[0] if entidad_row else str(entidad_nit))
                        prefijo = entidad_nombre_fs.strip().upper()[:3].replace(' ', '')
                        carpeta_entidad = f"{prefijo}_{entidad_nit}"
                        base_dir = os.path.join(os.path.dirname(__file__), 'reportes_liquidacion', carpeta_entidad)
                        os.makedirs(base_dir, exist_ok=True)
                        target_path = os.path.join(base_dir, merged_name)
                        if os.path.exists(target_path):
                            base, ext = os.path.splitext(target_path)
                            suf = 2
                            while os.path.exists(f"{base}_v{suf}{ext}"):
                                suf += 1
                            target_path = f"{base}_v{suf}{ext}"
                        with open(target_path, 'wb') as f:
                            f.write(merged_bytes)
                        st.info(f"También se guardó en: {target_path}")
                    except Exception as _e_save_merge:
                        st.warning(f"No se pudo guardar automáticamente el PDF combinado en disco: {_e_save_merge}")

                    st.download_button(
                        label="⬇️ Descargar Carta + Consolidado",
                        data=merged_bytes,
                        file_name=merged_name,
                        mime="application/pdf",
                        key="dl_merge_carta_consol"
                    )
                    st.success("PDF combinado listo: primero la carta, luego el consolidado.")
                except Exception as e:
                    st.error(f"No fue posible unir la carta y el consolidado: {e}")
                finally:
                    pass
            st.caption("Genera un ZIP por cada entidad, uno tras otro, con la misma estructura de carpetas por año.")

            corr_all = st.toggle(
                "Si ya existe la cuenta del mismo periodo, actualizar (conservar consecutivo) [todas las entidades]",
                value=True,
                key="pol_dup_corr_all"
            )

            # Listado de entidades para selección con check
            try:
                from sqlalchemy import text
                entidades_all = session.execute(text("SELECT nit, nombre FROM entidad ORDER BY nombre")).fetchall()
            except Exception:
                entidades_all = []

            opciones_all = [f"{e[1]} ({e[0]})" for e in entidades_all]
            mapa_opt_a_nit = {f"{e[1]} ({e[0]})": str(e[0]) for e in entidades_all}
            nit_a_nombre = {str(e[0]): e[1] for e in entidades_all}

            # Buscador para filtrar entidades por nombre o NIT
            filtro = st.text_input("Buscar entidades (nombre o NIT)", value=st.session_state.get('filter_ents', ''), key="filter_ents")
            if filtro:
                f = filtro.strip().lower()
                opciones_filtradas = [opt for opt in opciones_all if f in opt.lower()]
            else:
                opciones_filtradas = opciones_all

            # Derivar lista filtrada como pares (nit, nombre) para la grilla
            filtradas_pairs = [(mapa_opt_a_nit[opt], nit_a_nombre[mapa_opt_a_nit[opt]]) for opt in opciones_filtradas]

            # Acciones rápidas sobre el conjunto filtrado
            csel1, csel2 = st.columns(2)
            with csel1:
                if st.button("Seleccionar todas (filtradas)"):
                    st.session_state['sel_nits'] = [nit for nit, _ in filtradas_pairs]
            with csel2:
                if st.button("Limpiar selección"):
                    st.session_state['sel_nits'] = []

            # Renderizar grilla con checkbox por entidad
            import pandas as pd
            sel_nits_state = set(st.session_state.get('sel_nits', []))
            data_rows = []
            for nit, nombre in filtradas_pairs:
                data_rows.append({
                    'Seleccionar': nit in sel_nits_state,
                    'NIT': nit,
                    'Nombre': nombre,
                })
            df_ent = pd.DataFrame(data_rows)
            edited = st.data_editor(
                df_ent,
                hide_index=True,
                use_container_width=True,
                column_config={
                    'Seleccionar': st.column_config.CheckboxColumn('Seleccionar', help='Marcar para incluir en la generación'),
                    'NIT': st.column_config.TextColumn('NIT', disabled=True),
                    'Nombre': st.column_config.TextColumn('Nombre', disabled=True),
                },
                key='grid_ents'
            )
            # Persistir selección editada
            try:
                st.session_state['sel_nits'] = [str(n) for n in edited.loc[edited['Seleccionar'], 'NIT'].tolist()]
            except Exception:
                pass

            if st.button("🔄 Generar ZIP de todas las entidades", key="btn_all_entities_zip"):
                if not entidades_all:
                    st.warning("No se encontraron entidades en la base de datos.")
                else:
                    barra = st.progress(0)
                    # Filtrar por selección (si hay)
                    selected_nits = set(st.session_state.get('sel_nits', []))
                    if selected_nits:
                        entidades_target = [(nit, nom) for (nit, nom) in entidades_all if str(nit) in selected_nits]
                    else:
                        entidades_target = entidades_all

                    total_e = len(entidades_target)
                    resultados = []
                    for i, (nit_e, nom_e) in enumerate(entidades_target, start=1):
                        try:
                            todas_min = _construir_todas_cuentas_min(session, nit_e, fecha_corte)
                            zip_bytes = generar_zip_masivo_completo(
                                entidad_nit=str(nit_e),
                                entidad_nombre=nom_e,
                                todas_las_cuentas=todas_min,
                                fecha_corte=fecha_corte,
                                corregir_existentes=corr_all
                            )
                            # Persistir también en disco (misma carpeta que personalizados)
                            try:
                                import os
                                prefijo = nom_e.strip().upper()[:3].replace(' ', '')
                                carpeta_entidad = f"{prefijo}_{nit_e}"
                                base_dir = os.path.join(os.path.dirname(__file__), 'reportes_liquidacion', carpeta_entidad)
                                os.makedirs(base_dir, exist_ok=True)
                                file_name = f"LIQUIDACION_MASIVA_{nit_e}_{fecha_corte.strftime('%Y%m%d')}.zip"
                                target_path = os.path.join(base_dir, file_name)
                                if os.path.exists(target_path):
                                    basep, extp = os.path.splitext(target_path)
                                    suf = 2
                                    while os.path.exists(f"{basep}_v{suf}{extp}"):
                                        suf += 1
                                    target_path = f"{basep}_v{suf}{extp}"
                                with open(target_path, 'wb') as f:
                                    f.write(zip_bytes)
                            except Exception as _e:
                                target_path = None
                                st.warning(f"No se pudo guardar en disco para {nom_e} ({nit_e}): {_e}")

                            resultados.append((nit_e, nom_e, zip_bytes, target_path))
                        except Exception as ex:
                            st.error(f"Error generando ZIP para {nom_e} ({nit_e}): {ex}")
                        finally:
                            barra.progress(i/total_e)

                    st.success(f"Proceso finalizado. Se generaron {len(resultados)} ZIP(s). Descárgalos abajo:")
                    for nit_e, nom_e, data_e, path_e in resultados:
                        st.download_button(
                            label=f"⬇️ {nom_e} ({nit_e})",
                            data=data_e,
                            file_name=f"LIQUIDACION_MASIVA_{nit_e}_{fecha_corte.strftime('%Y%m%d')}.zip",
                            mime="application/zip",
                            key=f"dl_zip_{nit_e}"
                        )
                        if path_e:
                            st.caption(f"Guardado en: {path_e}")

            # Los botones persistentes ya se renderizan en los placeholders superiores
        
        session.close()
        
    except Exception as e:
        st.error(f"Error al conectar con la base de datos: {e}")
        import traceback
        st.error(traceback.format_exc())

# (secciones duplicadas eliminadas: mantener solo las versiones actualizadas más abajo)

    

# --- Módulo Seguridad y Trazabilidad ---
elif menu == "🔒 Seguridad y Trazabilidad":
    st.title("Seguridad y Trazabilidad")
    st.subheader("Registro de acciones y accesos por usuario")

    st.write("(Aquí se mostraría la bitácora de acciones y roles de usuario)")
    st.selectbox("Rol de usuario", ["Administrador", "Analista", "Auditor"])
    st.button("Ver historial de acciones")
    st.button("Configurar accesos y roles")

    st.markdown("---")
    st.subheader("🧹 Mantenimiento: Reiniciar consecutivo y borrar historial")
    st.warning("Acción sensible: esta operación elimina cuentas de cobro en base de datos y reinicia el consecutivo. Haz un respaldo antes.")

    from app.db import get_session
    from sqlalchemy import text
    session = get_session()
    try:
        entidades = session.execute(text("SELECT nit, nombre FROM entidad ORDER BY nombre")).fetchall()
    except Exception:
        entidades = []

    # Selección de entidades (vacío = todas)
    etiquetas = [f"{e[1]} ({e[0]})" for e in entidades]
    mapa = {f"{e[1]} ({e[0]})": str(e[0]) for e in entidades}
    filtro_txt = st.text_input("Buscar entidades (nombre o NIT)", key="admin_filter_ents")
    if filtro_txt:
        f = filtro_txt.strip().lower()
        opciones = [lab for lab in etiquetas if f in lab.lower()]
    else:
        opciones = etiquetas
    seleccion = st.multiselect("Selecciona entidades a afectar (vacío = todas)", options=opciones, key="admin_sel_ents")

    # Vista previa de conteos
    with st.expander("Ver conteos actuales en cuenta_cobro"):
        try:
            rows = session.execute(text("SELECT nit_entidad, COUNT(*) c FROM cuenta_cobro GROUP BY nit_entidad ORDER BY c DESC")).fetchall()
            if not rows:
                st.info("No hay registros en cuenta_cobro.")
            else:
                import pandas as pd
                df = pd.DataFrame(rows, columns=["NIT", "Cuentas"])
                if seleccion:
                    nits_sel = {mapa[s] for s in seleccion if s in mapa}
                    df = df[df["NIT"].astype(str).isin(nits_sel)]
                st.dataframe(df, use_container_width=True, height=250)
        except Exception as ex:
            st.error(f"No fue posible obtener conteos: {ex}")

    # Consulta rápida de consecutivos desde seguridad
    with st.expander("Consecutivos actuales (BD vs archivo)"):
        if st.button("🔄 Recalcular siguiente consecutivo", key="btn_recalc_cons_admin"):
            try:
                mx = session.execute(text("SELECT MAX(consecutivo) FROM cuenta_cobro")).scalar()
            except Exception:
                mx = None
            try:
                with open('ultimo_consecutivo.txt', 'r', encoding='utf-8') as f:
                    u = int(f.read().strip())
            except Exception:
                u = None
            cola, colb = st.columns(2)
            with cola:
                st.metric("Siguiente BD", (mx or 0) + 1)
                st.caption(f"MAX BD: {mx or 0}")
            with colb:
                st.metric("Siguiente archivo", (u or 0) + 1)
                st.caption(f"Archivo actual: {u if u is not None else 'N/A'}")

    colA, colB = st.columns(2)
    with colA:
        reset_txt_file = st.checkbox("Reiniciar archivo ultimo_consecutivo.txt a 0", value=True)
        reset_ai = st.checkbox("Reiniciar AUTO_INCREMENT de cuenta_cobro_id a 1", value=False)
        wipe_reportes = st.checkbox("Borrar archivos/ZIPs en carpeta reportes_liquidacion", value=True)
        wipe_temp = st.checkbox("Borrar carpeta temporal temp_pdf_generation", value=True)
    with colB:
        scope_all = st.checkbox("Borrar TODAS las entidades", value=(len(seleccion) == 0))
        st.caption("Si no seleccionas ninguna entidad, se asume TODAS.")
        wipe_hist_tables = st.checkbox("Borrar tablas de historial (liquidacion, liquidacion_detalle, periodo_liquidacion, pago)", value=True)

    confirm = st.text_input("Escribe BORRAR para confirmar", value="", key="confirm_borrar")

    if st.button("❌ Borrar historial y reiniciar consecutivo", type="primary"):
        if confirm.strip().upper() != "BORRAR":
            st.error("Debes escribir BORRAR para ejecutar esta acción.")
        else:
            try:
                # Determinar NITs a afectar
                if scope_all or not seleccion:
                    nits = [str(e[0]) for e in entidades]
                else:
                    nits = [mapa[s] for s in seleccion if s in mapa]

                # Borrar cuenta_cobro por cada NIT
                total_borradas = 0
                for nit in nits:
                    res = session.execute(text("DELETE FROM cuenta_cobro WHERE nit_entidad = :nit"), {"nit": nit})
                    total_borradas += res.rowcount if res.rowcount is not None else 0
                session.commit()

                # Borrar en tablas de historial relacionadas, respetando FKs y el esquema real
                if wipe_hist_tables:
                    borradas_tablas = {}

                    # Helper: verificar existencia de una tabla (para entornos donde no exista liquidacion_detalle)
                    def _tabla_existe(nombre_tbl: str) -> bool:
                        try:
                            return bool(session.execute(
                                text("SELECT COUNT(*) FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = :tbl"),
                                {"tbl": nombre_tbl}
                            ).scalar())
                        except Exception:
                            return False

                    # Para cada NIT, capturar los periodo_liquidacion_id referenciados por pagos ANTES de borrar pagos
                    periodos_por_nit = {}
                    try:
                        for nit in nits:
                            ids = session.execute(text(
                                "SELECT DISTINCT periodo_liquidacion_id FROM pago "
                                "WHERE periodo_liquidacion_id IS NOT NULL "
                                "AND pensionado_id IN (SELECT pensionado_id FROM pensionado WHERE nit_entidad = :nit)"
                            ), {"nit": nit}).fetchall()
                            # Almacenar como lista de int simples
                            ids_int = []
                            for row in ids:
                                val = row[0]
                                if val is not None:
                                    try:
                                        ids_int.append(int(val))
                                    except Exception:
                                        pass
                            periodos_por_nit[nit] = ids_int
                    except Exception as ex_pids:
                        st.warning(f"No se pudieron identificar periodos a eliminar: {ex_pids}")

                    # 1) liquidacion_detalle (si existe): tiene pensionado_id
                    if _tabla_existe("liquidacion_detalle"):
                        try:
                            borradas = 0
                            for nit in nits:
                                res = session.execute(text(
                                    "DELETE FROM liquidacion_detalle WHERE pensionado_id IN (SELECT pensionado_id FROM pensionado WHERE nit_entidad = :nit)"
                                ), {"nit": nit})
                                if res.rowcount is not None:
                                    borradas += res.rowcount
                            session.commit()
                            borradas_tablas["liquidacion_detalle"] = borradas
                        except Exception as ex_del:
                            session.rollback()
                            st.warning(f"No se pudo limpiar liquidacion_detalle: {ex_del}")

                    # 2) liquidacion: tiene pensionado_id
                    try:
                        borradas = 0
                        for nit in nits:
                            res = session.execute(text(
                                "DELETE FROM liquidacion WHERE pensionado_id IN (SELECT pensionado_id FROM pensionado WHERE nit_entidad = :nit)"
                            ), {"nit": nit})
                            if res.rowcount is not None:
                                borradas += res.rowcount
                        session.commit()
                        borradas_tablas["liquidacion"] = borradas
                    except Exception as ex_del:
                        session.rollback()
                        st.warning(f"No se pudo limpiar liquidacion: {ex_del}")

                    # 3) pago: tiene pensionado_id
                    try:
                        borradas = 0
                        for nit in nits:
                            res = session.execute(text(
                                "DELETE FROM pago WHERE pensionado_id IN (SELECT pensionado_id FROM pensionado WHERE nit_entidad = :nit)"
                            ), {"nit": nit})
                            if res.rowcount is not None:
                                borradas += res.rowcount
                        session.commit()
                        borradas_tablas["pago"] = borradas
                    except Exception as ex_del:
                        session.rollback()
                        st.warning(f"No se pudo limpiar pago: {ex_del}")

                    # 4) periodo_liquidacion: NO tiene pensionado_id; borrar por IDs capturados desde pago
                    try:
                        borradas = 0
                        for nit in nits:
                            ids = periodos_por_nit.get(nit, [])
                            if not ids:
                                continue
                            # Construir IN seguro con enteros ya leídos de BD
                            ids_str = ",".join(str(i) for i in ids)
                            res = session.execute(text(
                                f"DELETE FROM periodo_liquidacion WHERE periodo_liquidacion_id IN ({ids_str})"
                            ))
                            if res.rowcount is not None:
                                borradas += res.rowcount
                        session.commit()
                        borradas_tablas["periodo_liquidacion"] = borradas
                    except Exception as ex_del:
                        session.rollback()
                        st.warning(f"No se pudo limpiar periodo_liquidacion: {ex_del}")

                # Eliminar archivos de reportes si procede
                import os, shutil
                mensajes_files = []
                if wipe_reportes:
                    base_dir = os.path.join(os.path.dirname(__file__), 'reportes_liquidacion')
                    try:
                        if scope_all or not seleccion:
                            if os.path.exists(base_dir):
                                shutil.rmtree(base_dir, ignore_errors=True)
                                os.makedirs(base_dir, exist_ok=True)
                                mensajes_files.append("Se vació la carpeta reportes_liquidacion por completo.")
                            else:
                                mensajes_files.append("La carpeta reportes_liquidacion no existe.")
                        else:
                            # Borrar carpetas por entidad: <PREFIJO>_<NIT>
                            mapa_nit_nombre = {str(e[0]): e[1] for e in entidades}
                            for nit in nits:
                                nombre = mapa_nit_nombre.get(str(nit), str(nit))
                                prefijo = nombre.strip().upper()[:3].replace(' ', '')
                                carpeta = f"{prefijo}_{nit}"
                                ruta = os.path.join(base_dir, carpeta)
                                if os.path.exists(ruta):
                                    shutil.rmtree(ruta, ignore_errors=True)
                                    mensajes_files.append(f"Eliminada carpeta de reportes: {carpeta}")
                    except Exception as exf:
                        st.warning(f"No se pudieron eliminar archivos de reportes: {exf}")

                if wipe_temp:
                    try:
                        temp_dir = os.path.join(os.path.dirname(__file__), 'temp_pdf_generation')
                        if os.path.exists(temp_dir):
                            shutil.rmtree(temp_dir, ignore_errors=True)
                    except Exception as ext:
                        st.warning(f"No se pudo limpiar carpeta temporal: {ext}")

                # Reiniciar AUTO_INCREMENT si procede (solo estructura, no afecta consecutivo lógico)
                if reset_ai:
                    try:
                        session.execute(text("ALTER TABLE cuenta_cobro AUTO_INCREMENT = 1"))
                        session.commit()
                    except Exception as ex_ai:
                        st.warning(f"No se pudo reiniciar AUTO_INCREMENT: {ex_ai}")

                # Reiniciar archivo local de consecutivo (usado por PDF consolidado)
                if reset_txt_file:
                    try:
                        with open('ultimo_consecutivo.txt', 'w', encoding='utf-8') as f:
                            f.write('0')
                    except Exception as exf:
                        st.warning(f"No se pudo reiniciar ultimo_consecutivo.txt: {exf}")

                # Mensaje final
                detalles_hist = ""
                if wipe_hist_tables:
                    try:
                        # Mostrar sumario si lo calculamos
                        detalles_hist = " | " + ", ".join([f"{k}: {v}" for k, v in borradas_tablas.items()])
                    except Exception:
                        pass
                extra_files = (" | " + " ".join(mensajes_files)) if mensajes_files else ""
                st.success(f"Operación completada. cuenta_cobro borradas: {total_borradas}{detalles_hist}{extra_files}")
            except Exception as ex:
                session.rollback()
                st.error(f"Error durante la operación: {ex}")
    session.close()

# --- Módulo Liquidar por Periodos Personalizados ---
elif menu == "🗓️ Liquidar por Periodos Personalizados":
    st.title("Liquidar por Periodos Personalizados")
    st.write("Selecciona la entidad, el rango de fechas y los meses que deseas excluir. Se generará un PDF individual por cada periodo seleccionado.")

    from datetime import date
    from app.db import get_session
    session = get_session()
    entidades = session.execute(text("SELECT nit, nombre FROM entidad ORDER BY nombre")).fetchall()
    entidad_nit = st.selectbox("Entidad", options=[e[0] for e in entidades], format_func=lambda x: next(e[1] for e in entidades if e[0]==x))

    meses = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"]

    modo = st.radio("Modo de selección", ["Rango de fechas", "Un solo mes"], horizontal=True)

    fecha_inicio = None
    fecha_fin = None
    excluir_meses = []
    solo_un_mes = True
    excluir_meses_por_año = {}

    if modo == "Rango de fechas":
        col1, col2 = st.columns(2)
        with col1:
            # Inicio no puede ser antes de Mar 2023 por la ventana de 30 meses definida
            fecha_inicio = st.date_input("Fecha inicio", value=date(2023,3,1), min_value=date(2023,3,1), max_value=date(2025,8,31))
        with col2:
            fecha_fin = st.date_input("Fecha fin", value=date(2025,8,31), min_value=fecha_inicio or date(2023,3,1), max_value=date(2025,8,31))
    # Excluir meses por año dentro del rango seleccionado
        st.markdown("**Excluir meses por año**")
        # Botón para limpiar exclusiones rápidamente (estabiliza estado al cambiar de ventana/rango)
        if st.button("Limpiar exclusiones", key="btn_limpiar_exclusiones"):
            for a in range(2022, 2026):
                k = f"excluir_{a}"
                if k in st.session_state:
                    del st.session_state[k]
        global_inicio = date(2022, 9, 1)
        global_fin = date(2025, 8, 31)
        # Iterar por cada año del rango seleccionado
        # Renderizar SIEMPRE los años de 2022 a 2025 para evitar cambios dinámicos en el DOM
        for año in range(2022, 2026):
            # Mantener opciones constantes (12 meses) para estabilidad del DOM
            keyname = f"excluir_{año}"
            # Sanear posibles valores previos fuera de las opciones válidas
            if keyname in st.session_state:
                prev = st.session_state[keyname]
                if isinstance(prev, (list, tuple)):
                    filtrado = [v for v in prev if v in meses]
                    if filtrado != prev:
                        st.session_state[keyname] = filtrado
            seleccion = st.multiselect(
                f"Excluir meses {año}",
                options=meses,
                help="Los meses fuera del rango seleccionado se ignoran automáticamente",
                key=keyname
            )
            # Mapear nombres seleccionados a números de mes
            excluir_meses_por_año[año] = set(meses.index(nombre) + 1 for nombre in seleccion)
        solo_un_mes = st.checkbox("Generar solo ese mes por PDF (no acumulado)", value=True)
    else:
        # Un solo mes: permitir elegir año y mes dentro de [Sep 2022, Ago 2025]
        col1, col2 = st.columns(2)
        with col1:
            año_sel = st.selectbox("Año", options=[2022, 2023, 2024, 2025], index=0)
        with col2:
            # Limitar meses según año seleccionado
            if año_sel == 2022:
                meses_idx = list(range(9, 13))  # Sep-Dic
            elif año_sel == 2025:
                meses_idx = list(range(1, 9))   # Ene-Ago
            else:
                meses_idx = list(range(1, 13))  # Todo el año
            nombres_filtrados = [meses[i-1] for i in meses_idx]
            mes_nombre = st.selectbox("Mes", options=nombres_filtrados)
            mes_sel = meses.index(mes_nombre) + 1

        # Construir fechas de inicio y fin del mes seleccionado
        from calendar import monthrange
        fecha_inicio = date(año_sel, mes_sel, 1)
        dia_fin = monthrange(año_sel, mes_sel)[1]
        fecha_fin = date(año_sel, mes_sel, dia_fin)
        solo_un_mes = True  # Forzar modo de un solo mes

    # PREVISUALIZACIÓN: construir lista de periodos resultantes antes de generar
    periodos_preview = []
    if modo == "Un solo mes" and fecha_inicio is not None:
        periodos_preview.append((fecha_inicio.year, fecha_inicio.month))
    elif modo == "Rango de fechas" and fecha_inicio is not None and fecha_fin is not None:
        actual = fecha_inicio
        while actual <= fecha_fin:
            excl_for_year = excluir_meses_por_año.get(actual.year, set())
            if actual.month not in excl_for_year:
                periodos_preview.append((actual.year, actual.month))
            if actual.month == 12:
                actual = date(actual.year+1, 1, 1)
            else:
                actual = date(actual.year, actual.month+1, 1)

    # Guardar en session_state para que métricas y generación usen exactamente la misma lista
    st.session_state['periodos_preview'] = periodos_preview

    # Contar pensionados de la entidad para estimar PDFs
    try:
        total_pensionados = session.execute(
            text("SELECT COUNT(*) FROM pensionado WHERE nit_entidad = :nit"), {"nit": entidad_nit}
        ).scalar() or 0
    except Exception:
        total_pensionados = 0

    colp1, colp2, colp3 = st.columns(3)
    with colp1:
        st.metric("Periodos seleccionados", len(st.session_state.get('periodos_preview', [])))
    with colp2:
        st.metric("Pensionados", total_pensionados)
    with colp3:
        st.metric("PDFs estimados", len(st.session_state.get('periodos_preview', [])) * total_pensionados)

    # Lista legible de periodos
    if st.session_state.get('periodos_preview'):
        etiquetas = [f"{meses[m-1]} {a}" for a, m in st.session_state['periodos_preview']]
        with st.expander("Ver periodos seleccionados"):
            st.write(", ".join(etiquetas))

    # Política de duplicados para esta sección
    corr_custom = st.toggle(
        "Si ya existe la cuenta del mismo periodo, actualizar (conservar consecutivo)",
        value=True,
        help="Si está activo: corrige la cuenta existente conservando el consecutivo. Si está desactivado: crea una nueva con nuevo consecutivo.",
        key="pol_dup_corr_custom"
    )

    if st.button("Generar PDFs por periodo"):
        import generar_pdf_oficial as gpo
        from generar_pdf_oficial import generar_pdf_para_pensionado
        # Aplicar política de duplicados a nivel de generador
        try:
            gpo.CONSEC_CORRECCION = bool(corr_custom)
            gpo.CONSEC_OVERRIDE = None
        except Exception:
            pass
        pensionados = session.execute(
            text("SELECT identificacion, nombre, numero_mesadas, fecha_ingreso_nomina, empresa, base_calculo_cuota_parte, porcentaje_cuota_parte, nit_entidad FROM pensionado WHERE nit_entidad = :nit"),
            {"nit": entidad_nit}
        ).fetchall()
        # Construir carpeta base por entidad: <3 primeras letras>_<NIT>
        # Obtener nombre de entidad
        entidad_row = session.execute(text("SELECT nombre FROM entidad WHERE nit = :nit"), {"nit": entidad_nit}).fetchone()
        entidad_nombre = (entidad_row[0] if entidad_row else str(entidad_nit))
        prefijo = entidad_nombre.strip().upper()[:3].replace(' ', '')
        carpeta_entidad = f"{prefijo}_{entidad_nit}"
        base_dir = os.path.join(os.path.dirname(__file__), 'reportes_liquidacion', carpeta_entidad)
        os.makedirs(base_dir, exist_ok=True)
        # Pre-escaneo de duplicados para advertir
        try:
            periodo_fin_chk = date(2025, 8, 1)
            dups = 0
            for p in pensionados:
                ced = str(p[0])
                for a, m in st.session_state.get('periodos_preview', []):
                    ini = date(a, m, 1)
                    cnt = session.execute(
                        text("SELECT COUNT(*) FROM cuenta_cobro WHERE nit_entidad=:n AND pensionado_identificacion=:id AND periodo_inicio=:ini AND periodo_fin=:fin"),
                        {"n": str(entidad_nit), "id": ced, "ini": ini, "fin": periodo_fin_chk}
                    ).scalar() or 0
                    if cnt:
                        dups += 1
                        break
            if dups:
                st.warning(f"Se encontraron cuentas existentes para {dups} pensionado(s) en los periodos seleccionados. Se {'corregirán' if corr_custom else 'crearán nuevas'} según la opción elegida.")
            else:
                st.info("No se encontraron cuentas existentes para los periodos seleccionados.")
        except Exception:
            pass

        # Reusar la lista de periodos ya calculada en la previsualización
        for pensionado in pensionados:
            for año, mes in st.session_state.get('periodos_preview', []):
                generar_pdf_para_pensionado(pensionado, 'custom', año, mes, solo_mes=solo_un_mes, output_dir=base_dir)
        st.success(f"PDFs generados para {len(pensionados)} pensionados y {len(st.session_state.get('periodos_preview', []))} periodos seleccionados.")
        # Acceso rápido al dashboard tras la generación personalizada
        if st.button("📊 Ver Dashboard con resumen", key="btn_goto_dashboard_custom"):
            st.session_state["menu_principal"] = "🏠 Dashboard"
            st.rerun()